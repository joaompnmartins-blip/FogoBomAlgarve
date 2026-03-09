import sys
import os
import base64
from pathlib import Path
import json
from datetime import datetime, timedelta
import calendar as cal
import hashlib
import secrets

BASE_DIR = Path(__file__).resolve().parent
MEDIA_DIR = BASE_DIR / "media" / "preplan_photos"
MEDIA_DIR.mkdir(parents=True, exist_ok=True)
BP_MEDIA_DIR = BASE_DIR / "media" / "bp_photos"
BP_MEDIA_DIR.mkdir(parents=True, exist_ok=True)
LEAFLET_PATH = BASE_DIR / "leaflet_map.html"
WEATHER_MAP_PATH = BASE_DIR / "leaflet_weather_map.html"

# Carregar mapa de parcelas
html_bytes = LEAFLET_PATH.read_bytes()
encoded = base64.b64encode(html_bytes).decode("utf-8")

MAP_IFRAME = f"""
<iframe
  id="map-iframe"
  src="data:text/html;base64,{encoded}"
  style="width:100%; height:500px; border:none;"
></iframe>
"""
REGISTRY_MAP_PATH = BASE_DIR / "leaflet_registry_map.html"
registry_html_bytes = REGISTRY_MAP_PATH.read_bytes()
registry_encoded = base64.b64encode(registry_html_bytes).decode("utf-8")

REGISTRY_MAP_IFRAME = f"""
<iframe
  id="registry-map-iframe"
  src="data:text/html;base64,{registry_encoded}"
  style="width:100%; height:500px; border:none;"
></iframe>
"""

# EDIT MAP - for "Editar/Eliminar Parcelas" tab (full editing with zoom to parcel)
EDIT_MAP_PATH = BASE_DIR / "leaflet_edit_map.html"
edit_html_bytes = EDIT_MAP_PATH.read_bytes()
edit_encoded = base64.b64encode(edit_html_bytes).decode("utf-8")

EDIT_MAP_IFRAME = f"""
<iframe
  id="edit-map-iframe"
  src="data:text/html;base64,{edit_encoded}"
  style="width:100%; height:500px; border:none;"
></iframe>
"""
# Carregar mapa meteorológico (FiredPT) com injeção lazy de parcelas da BD
# A leitura do HTML é feita uma vez; a injeção das parcelas é feita em runtime
# (após o Django estar inicializado) pela função build_weather_map_iframe().
try:
    _FIREDPT_HTML_TEMPLATE = WEATHER_MAP_PATH.read_text(encoding="utf-8")
    _FIREDPT_AVAILABLE = True
except FileNotFoundError:
    _FIREDPT_HTML_TEMPLATE = ""
    _FIREDPT_AVAILABLE = False

def build_weather_map_iframe() -> str:
    """
    Constrói o iframe do mapa meteorológico FiredPT injetando as parcelas
    da base de dados como variável JavaScript pré-carregada.

    Chamada após o Django estar inicializado (dentro de create_main_interface).
    Usa o centróide da geometria de cada parcela como ponto representativo.
    """
    if not _FIREDPT_AVAILABLE:
        return """
        <div style="padding:20px; background:#1a1f1c; border-radius:8px; text-align:center; color:#9ab5a2;">
            <h3 style="color:#f5d84a;">⚠️ FiredPT não disponível</h3>
            <p>O arquivo leaflet_weather_map.html não foi encontrado.</p>
        </div>
        """

    # ── Recolher parcelas com centróide ──────────────────────────────────────
    db_parcels_js = "[]"
    try:
        from parcels.models import FireParcel
        parcels_qs = FireParcel.objects.all()
        parcels_list = []
        for p in parcels_qs:
            try:
                centroid = p.geometry.centroid
                parcels_list.append({
                    "id":   str(p.id),
                    "name": p.name,
                    "lat":  round(centroid.y, 6),
                    "lng":  round(centroid.x, 6),
                })
            except Exception:
                continue  # ignorar parcelas sem geometria válida
        db_parcels_js = json.dumps(parcels_list, ensure_ascii=False)
    except Exception as e:
        print(f"[FiredPT] Aviso: não foi possível carregar parcelas — {e}")

    # ── Substituir placeholder no template ───────────────────────────────────
    html = _FIREDPT_HTML_TEMPLATE.replace(
        "let parcels       = FIREDPT_DB_PARCELS;",
        f"let parcels       = {db_parcels_js};",
        1,
    )

    # ── Codificar em base64 e devolver iframe ─────────────────────────────────
    encoded = base64.b64encode(html.encode("utf-8")).decode("utf-8")
    return f"""
    <iframe
      id="weather-map-iframe"
      src="data:text/html;base64,{encoded}"
      style="width:100%; height:calc(100vh - 120px); min-height:600px; border:none; border-radius:8px;"
    ></iframe>
    """

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "backend"))
sys.path.insert(0, BASE_DIR)

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "fire_mgmt.settings")

import django
django.setup()

import gradio as gr
import pandas as pd
from parcels.models import FireParcel
from operatives.models import Operative
from fire_actions.models import FireAction, BurningPlan
from django.contrib.gis.geos import GEOSGeometry
from django.db.models import Count, Q
from django.contrib.auth.models import User
from django.contrib.auth import authenticate
import requests

# ============= AUTHENTICATION FUNCTIONS =============
class SessionManager:
    """Gestor de sessões de utilizadores"""
    def __init__(self):
        self.sessions = {}
    
    def create_session(self, username):
        """Criar nova sessão"""
        token = secrets.token_urlsafe(32)
        self.sessions[token] = {
            'username': username,
            'created_at': datetime.now(),
            'last_activity': datetime.now()
        }
        return token
    
    def validate_session(self, token):
        """Validar se sessão existe e está ativa"""
        if not token or token not in self.sessions:
            return False
        
        session = self.sessions[token]
        # Sessão expira após 8 horas de inatividade
        if datetime.now() - session['last_activity'] > timedelta(hours=8):
            del self.sessions[token]
            return False
        
        # Atualizar última atividade
        session['last_activity'] = datetime.now()
        return True
    
    def get_username(self, token):
        """Obter username da sessão"""
        if self.validate_session(token):
            return self.sessions[token]['username']
        return None
    
    def logout(self, token):
        """Terminar sessão"""
        if token in self.sessions:
            del self.sessions[token]

# Instância global do gestor de sessões
session_manager = SessionManager()

def login_user(username, password):
    """Autenticar utilizador"""
    if not username or not password:
        return None, "❌ Por favor preencha todos os campos"
    
    # Autenticar com Django
    user = authenticate(username=username, password=password)
    
    if user is not None:
        # Criar sessão
        token = session_manager.create_session(username)
        return token, f"✓ Bem-vindo, {username}!"
    else:
        return None, "❌ Credenciais inválidas"

def logout_user(token):
    """Terminar sessão do utilizador"""
    session_manager.logout(token)
    return None, "Sessão terminada. Por favor faça login novamente."

def register_user(username, email, password, confirm_password):
    """Registar novo utilizador"""
    if not all([username, email, password, confirm_password]):
        return "❌ Por favor preencha todos os campos"
    
    if password != confirm_password:
        return "❌ As passwords não coincidem"
    
    if len(password) < 6:
        return "❌ A password deve ter pelo menos 6 caracteres"
    
    if User.objects.filter(username=username).exists():
        return "❌ Nome de utilizador já existe"
    
    if User.objects.filter(email=email).exists():
        return "❌ Email já registado"
    
    try:
        User.objects.create_user(username=username, email=email, password=password)
        return f"✓ Utilizador '{username}' criado com sucesso! Pode fazer login agora."
    except Exception as e:
        return f"❌ Erro ao criar utilizador: {str(e)}"

def check_authentication(token):
    """Verificar se utilizador está autenticado"""
    if session_manager.validate_session(token):
        username = session_manager.get_username(token)
        return True, f"👤 {username}"
    return False, "🔒 Não autenticado"

# ============= ALGARVE CONCELHOS =============
ALGARVE_CONCELHOS = [
    "Albufeira", "Alcoutim", "Aljezur", "Castro Marim", "Faro",
    "Lagoa", "Lagos", "Loulé", "Monchique", "Olhão",
    "Portimão", "São Brás de Alportel", "Silves", "Tavira",
    "Vila do Bispo", "Vila Real de Santo António"
]

# ============= PARCELS FUNCTIONS (com autenticação) =============
def save_parcel(name, resp_name, resp_email, veg_type, infra, owner, concelho, geometry_json, area_ha):
    try:
        if not geometry_json or not geometry_json.strip():
            return "❌ Por favor desenhe a geometria da parcela no mapa antes de guardar.", get_parcels_geojson(), "", name, resp_name, resp_email, veg_type, infra, owner, concelho, geometry_json, area_ha

        try:
            final_area = float(area_ha) if area_ha and str(area_ha).strip() else None
        except (ValueError, TypeError):
            final_area = None

        geometry_dict = json.loads(geometry_json)
        geom = GEOSGeometry(json.dumps(geometry_dict), srid=4326)
        
        FireParcel.objects.create(
            name=name,
            resp_name=resp_name,
            resp_email=resp_email,
            vegetation_type=veg_type,
            infrastructure=infra,
            owner_info=owner,
            concelho=concelho,
            area_ha=final_area,
            geometry=geom
        )
        
        geojson = get_parcels_geojson()
        area_str = f", {final_area} ha" if final_area else ""
        concelho_str = f", {concelho}" if concelho else ""
        return (
            f"✓ Parcela '{name}' guardada!{area_str}{concelho_str}",
            geojson,
            "save",
            "",                    # p_name
            "",                    # p_resp
            "",                    # p_email
            "Matos",               # p_veg
            "Sem informação",      # p_infra
            "Sim",                 # p_owner
            "",              # p_concelho
            "",                    # geometry
            ""                     # area_ha
        )
    except Exception as e:
        return f"Erro: {str(e)}", get_parcels_geojson(), "", name, resp_name, resp_email, veg_type, infra, owner, concelho, geometry_json, area_ha

def clear_registry_form():
    """Clear the registry form and any unsaved drawn parcels"""
    return (
        "",           # p_name
        "",           # p_resp
        "",           # p_email
        "Matos",      # p_veg
        "Sem informação", # p_infra
        "Sim",        # p_owner
        "",              # p_concelho
        "",           # geometry
        "",           # output message
        get_parcels_geojson(),  # geojson (unchanged)
        "clear",      # Action flag
        ""            # area_ha
    )
    
def get_parcels_list():
    """
    Obtém a lista de parcelas com os novos campos para exibição na tabela Gradio.
    """
    # Incluímos os novos campos no .values()
    parcels = FireParcel.objects.all().values(
        'id', 'name', 'resp_name', 'vegetation_type', 'infrastructure', 'owner_info', 'concelho'
    )
    
    # Retornamos uma lista de listas (formato esperado pelo gr.Dataframe)
    return [
        [
            p['id'], 
            p['name'], 
            p['resp_name'], 
            p['vegetation_type'], 
            p['infrastructure'], 
            p['owner_info'],
            p.get('concelho', '')
        ] for p in parcels
    ]

def get_parcels_geojson():
    """
    Gera o GeoJSON para o mapa, incluindo os metadados dos novos campos nas propriedades.
    """
    parcels = FireParcel.objects.all()
    features = []
    for parcel in parcels:
        try:
            # Converte a geometria do Django para dicionário Python
            geometry = json.loads(parcel.geometry.geojson)
            
            features.append({
                "type": "Feature",
                "properties": {
                    "id": parcel.id,
                    "name": parcel.name,
                    "responsavel": parcel.resp_name,
                    "email": parcel.resp_email,
                    "vegetation_type": parcel.vegetation_type,
                    "infraestrutura": parcel.infrastructure,
                    "tem_proprietario": parcel.owner_info,
                    "concelho": getattr(parcel, 'concelho', ''),
                    
                },
                "geometry": geometry
            })
        except Exception as e:
            print(f"Erro ao processar parcela {parcel.id}: {e}")
            continue
            
    return json.dumps({"type": "FeatureCollection", "features": features})

def reload_map_parcels():
    geojson = get_parcels_geojson()
    return f"""
    <script>
    (function() {{
        const iframe = document.getElementById('map-iframe');
        if (iframe && iframe.contentWindow) {{
            iframe.contentWindow.postMessage({{
                type: 'load_parcels',
                geojson: {geojson}
            }}, '*');
        }}
    }})();
    </script>
    """
def load_parcel_for_edit(evt: gr.SelectData, current_table_data):
    try:
        row_index = evt.index[0]
        parcel_id = current_table_data.iloc[row_index, 0] 
        
        p = FireParcel.objects.get(id=int(parcel_id))
        
        # CORREÇÃO AQUI: 
        # No GeoDjango usa-se p.geometry.json para obter a string GeoJSON
        geom_geojson = p.geometry.json 
        
        return (
            p.id, 
            p.name, 
            p.resp_name, 
            p.resp_email, 
            p.vegetation_type, 
            p.infrastructure, 
            p.owner_info,
            geom_geojson, # Agora retorna a string JSON correta
            f"✅ Parcela '{p.name}' pronta para edição."
        )
    except Exception as e:
        print(f"Erro detalhado: {e}")
        return None, "", "", "", "Matos", "Sem informação", "Sim", "", f"❌ Erro: {str(e)}"

def update_parcel(p_id, name, resp, email, veg, infra, owner, geom_json):
    """Atualiza os dados na base de dados"""
    if not p_id: return "❌ Selecione primeiro uma parcela!", get_parcels_list(), ""
    try:
        p = FireParcel.objects.get(id=p_id)
        p.name, p.resp_name, p.resp_email = name, resp, email
        p.vegetation_type, p.infrastructure, p.owner_info = veg, infra, owner
        if geom_json:
            p.geometry = GEOSGeometry(json.dumps(json.loads(geom_json)))
        p.save()
        return f"✅ Parcela '{name}' atualizada!", get_parcels_list(), reload_map_parcels()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_parcels_list(), ""

def delete_parcel(p_id):
    """Elimina a parcela"""
    if not p_id: return "❌ Selecione uma parcela!", get_parcels_list(), ""
    try:
        p = FireParcel.objects.get(id=p_id)
        name = p.name
        p.delete()
        return f"🗑️ Parcela '{name}' eliminada!", get_parcels_list(), reload_map_parcels()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_parcels_list(), ""
    
# ============= PARCEL EDIT FUNCTIONS FOR DROPDOWN =============

def get_parcel_by_id(parcel_id):
    """
    Get a single parcel by ID for editing
    Returns a dictionary with parcel data or None if not found
    """
    try:
        parcel = FireParcel.objects.get(id=parcel_id)
        
        raw_owner = parcel.owner_info or ""
        if raw_owner.lower() in ("yes", "sim"):
            owner_norm = "Sim"
        elif raw_owner.lower() in ("no", "não", "nao"):
            owner_norm = "Não"
        else:
            owner_norm = raw_owner

        return {
            'id': parcel.id,
            'name': parcel.name,
            'responsible': parcel.resp_name,
            'email': parcel.resp_email,
            'vegetation': parcel.vegetation_type,
            'infrastructure': parcel.infrastructure,
            'owner_info': owner_norm,
            'concelho': getattr(parcel, 'concelho', ''),
            'geometry': parcel.geometry.json,  # GeoJSON string
            'area_ha': getattr(parcel, 'area_ha', None)
        }
    except FireParcel.DoesNotExist:
        return None
    except Exception as e:
        print(f"Error loading parcel {parcel_id}: {e}")
        return None


def load_parcel_for_edit_dropdown(parcel_id):
    """
    Load parcel data when selected from dropdown
    Returns tuple with all form fields populated
    """
    if not parcel_id:
        return (
            None,  # edit_id
            "",    # e_name
            "",    # e_resp
            "",    # e_email
            None,  # e_veg
            None,  # e_infra
            None,  # e_owner
            "",    # e_concelho
            "",    # e_area
            "",    # edit_geom
            ""     # edit_msg
        )
    
    parcel = get_parcel_by_id(parcel_id)
    
    if parcel:
        return (
            parcel['id'],
            parcel['name'],
            parcel['responsible'],
            parcel['email'],
            parcel['vegetation'],
            parcel['infrastructure'],
            parcel['owner_info'],
            parcel['concelho'] or '',
            str(parcel['area_ha']) if parcel['area_ha'] is not None else '',
            parcel['geometry'],
            f"✅ Parcela '{parcel['name']}' carregada para edição"
        )
    else:
        return (
            None,
            "",
            "",
            "",
            None,
            None,
            None,
            "",    # e_concelho
            "",    # e_area
            "",
            "❌ Erro: Parcela não encontrada"
        )


def update_parcel_dropdown(parcel_id, name, resp, email, veg, infra, owner, concelho, area_ha, geom_json):
    """Update parcel and reload EDIT map"""
    if not parcel_id:
        return (
            "❌ Selecione uma parcela primeiro!",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
    
    if not name or not resp:
        return (
            "❌ Nome e Responsável são obrigatórios!",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
    
    try:
        parcel = FireParcel.objects.get(id=parcel_id)
        parcel.name = name
        parcel.resp_name = resp
        parcel.resp_email = email
        parcel.vegetation_type = veg
        parcel.infrastructure = infra
        parcel.owner_info = owner
        parcel.concelho = concelho
        try:
            parcel.area_ha = float(area_ha) if area_ha and str(area_ha).strip() else parcel.area_ha
        except (ValueError, TypeError):
            pass
        
        if geom_json and geom_json.strip():
            try:
                parcel.geometry = GEOSGeometry(json.dumps(json.loads(geom_json)))
            except Exception as e:
                return (
                    f"❌ Erro na geometria: {str(e)}",
                    gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
                    ""
                )
        
        parcel.save()
        
        new_choices = [(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]
        
        return (
            f"✅ Parcela '{name}' atualizada com sucesso!",
            gr.Dropdown(choices=new_choices, value=None),  # Reset dropdown
            reload_edit_map(),  # Reload the EDIT map
            None,   # edit_id
            "",     # e_name
            "",     # e_resp
            "",     # e_email
            None,   # e_veg
            None,   # e_infra
            None,   # e_owner
            "",     # e_concelho
            "",     # e_area
            ""      # edit_geom
        )
    except FireParcel.DoesNotExist:
        return (
            "❌ Parcela não encontrada",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
    except Exception as e:
        return (
            f"❌ Erro ao atualizar: {str(e)}",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )


def delete_parcel_dropdown(parcel_id):
    """Delete parcel and reload EDIT map"""
    if not parcel_id:
        return (
            "❌ Selecione uma parcela primeiro!",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
    
    try:
        parcel = FireParcel.objects.get(id=parcel_id)
        name = parcel.name
        
        actions_count = parcel.fire_actions.count() if hasattr(parcel, 'fire_actions') else 0
        if actions_count > 0:
            return (
                f"⚠️ Não é possível eliminar '{name}' - tem {actions_count} ação(ões) associada(s)!",
                gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
                ""
            )
        
        parcel.delete()
        
        new_choices = [(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]
        
        return (
            f"🗑️ Parcela '{name}' eliminada com sucesso!",
            gr.Dropdown(choices=new_choices, value=None),
            reload_edit_map()  # Reload the EDIT map
        )
    except FireParcel.DoesNotExist:
        return (
            "❌ Parcela não encontrada",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
    except Exception as e:
        return (
            f"❌ Erro ao eliminar: {str(e)}",
            gr.Dropdown(choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()]),
            ""
        )
def zoom_to_parcel(parcel_id):
    """
    Returns the parcel GeoJSON geometry string so the JS .then() chain
    can send it directly to the iframe via postMessage.
    (Injecting <script> tags via gr.HTML is sandboxed by browsers and never executes.)
    """
    if not parcel_id:
        return ""
    try:
        p = FireParcel.objects.get(id=int(parcel_id))
        return p.geometry.json
    except Exception:
        return ""
def reload_registry_map():
    """Reload parcels on the REGISTRY map"""
    geojson = get_parcels_geojson()
    return f"""
    <script>
    (function() {{
        const iframe = document.getElementById('registry-map-iframe');
        if (iframe && iframe.contentWindow) {{
            iframe.contentWindow.postMessage({{
                type: 'load_parcels',
                geojson: {geojson}
            }}, '*');
        }}
    }})();
    </script>
    """

def reload_edit_map():
    """Reload parcels on the EDIT map"""
    geojson = get_parcels_geojson()
    return f"""
    <script>
    (function() {{
        const iframe = document.getElementById('edit-map-iframe');
        if (iframe && iframe.contentWindow) {{
            iframe.contentWindow.postMessage({{
                type: 'load_parcels',
                geojson: {geojson}
            }}, '*');
        }}
    }})();
    </script>
    """
def get_filtered_parcels(concelho_filter, resp_filter, veg_filter):
    """Filter parcels by concelho, responsavel and vegetation type"""
    parcels = FireParcel.objects.all()
    
    if concelho_filter and concelho_filter != "Todos":
        parcels = parcels.filter(concelho=concelho_filter)
    
    if resp_filter and resp_filter.strip():
        parcels = parcels.filter(resp_name__icontains=resp_filter.strip())
    
    if veg_filter and veg_filter != "Todos":
        parcels = parcels.filter(vegetation_type=veg_filter)
    
    return [
        (f"{p.name} ({p.concelho or 'sem concelho'})", p.id)
        for p in parcels
    ]

# ============= OPERATIVES FUNCTIONS (com autenticação) =============
def get_operatives_stats():
    """Stats bar: total + breakdown by certification."""
    operatives = Operative.objects.all()
    total = operatives.count()
    by_cert = operatives.values('certification_level').annotate(count=Count('certification_level'))
    cert_map = {
        "Observador": "👁️",
        "Operacional Queima": "🔥",
        "Tecnico Fogo Controlado": "🧯",
        "Outro": "👤",
    }
    lines = [f"**👥 Total: {total}**"]
    for item in by_cert:
        cert = item['certification_level'] or "—"
        icon = cert_map.get(cert, "👤")
        lines.append(f"{icon} {cert}: **{item['count']}**")
    return "  |  ".join(lines)

def get_filtered_operatives(search_term, certification_filter):
    operatives = Operative.objects.all()
    if search_term and search_term.strip():
        operatives = operatives.filter(
            Q(name__icontains=search_term) |
            Q(email__icontains=search_term) |
            Q(phone__icontains=search_term) |
            Q(nif__icontains=search_term)
        )
    if certification_filter and certification_filter != "Todos":
        operatives = operatives.filter(certification_level=certification_filter)
    result = []
    for op in operatives:
        result.append([
            op.id,
            op.name,
            getattr(op, 'nif', '') or '',
            op.email,
            op.phone or '',
            op.certification_level,
            op.burning_plans.count()
        ])
    return result

def load_operative_for_edit(op_id):
    """Load a single operative by ID into the edit form."""
    if not op_id:
        return None, "", "", "", "", "Observador", ""
    try:
        op = Operative.objects.get(id=int(op_id))
        return (
            op.id,
            op.name,
            getattr(op, 'nif', '') or '',
            op.email,
            op.phone or '',
            op.certification_level,
            getattr(op, 'notes', '') or ''
        )
    except Operative.DoesNotExist:
        return None, "", "", "", "", "Observador", ""

def save_operative(op_id, name, nif, email, phone, certification, notes):
    """Create or update an operative."""
    if not name or not name.strip():
        return "❌ O nome é obrigatório.", get_filtered_operatives("", ""), get_operatives_stats()
    if not email or not email.strip():
        return "❌ O email é obrigatório.", get_filtered_operatives("", ""), get_operatives_stats()

    try:
        if op_id:
            # UPDATE
            op = Operative.objects.get(id=int(op_id))
            op.name = name.strip()
            try: op.nif = nif.strip() if nif else ""
            except: pass
            op.email = email.strip()
            op.phone = phone.strip() if phone else ""
            op.certification_level = certification
            try: op.notes = notes.strip() if notes else ""
            except: pass
            op.save()
            msg = f"✅ Operacional '{name}' atualizado com sucesso!"
        else:
            # CREATE — check duplicate email
            if Operative.objects.filter(email__iexact=email.strip()).exists():
                return f"❌ Já existe um operacional com o email '{email}'.", get_filtered_operatives("", ""), get_operatives_stats()
            kwargs = dict(name=name.strip(), email=email.strip(),
                          phone=phone.strip() if phone else "",
                          certification_level=certification)
            try: kwargs['nif'] = nif.strip() if nif else ""
            except: pass
            try: kwargs['notes'] = notes.strip() if notes else ""
            except: pass
            Operative.objects.create(**kwargs)
            msg = f"✅ Operacional '{name}' criado com sucesso!"

        return msg, get_filtered_operatives("", ""), get_operatives_stats()
    except Operative.DoesNotExist:
        return "❌ Operacional não encontrado.", get_filtered_operatives("", ""), get_operatives_stats()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_filtered_operatives("", ""), get_operatives_stats()

def delete_operative_new(op_id):
    """Delete an operative if it has no associated actions."""
    if not op_id:
        return "❌ Selecione um operacional primeiro.", get_filtered_operatives("", ""), get_operatives_stats()
    try:
        op = Operative.objects.get(id=int(op_id))
        name = op.name
        actions_count = op.burning_plans.count()
        if actions_count > 0:
            return (f"⚠️ Não é possível eliminar '{name}' — tem {actions_count} Plano(s) de Queima associado(s).",
                    get_filtered_operatives("", ""), get_operatives_stats())
        op.delete()
        return f"✅ Operacional '{name}' eliminado.", get_filtered_operatives("", ""), get_operatives_stats()
    except Operative.DoesNotExist:
        return "❌ Operacional não encontrado.", get_filtered_operatives("", ""), get_operatives_stats()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_filtered_operatives("", ""), get_operatives_stats()

# Keep legacy stubs so fire-actions tab references don't break
def create_operative(name, email, phone, certification):
    msg, _, stats = save_operative(None, name, "", email, phone, certification, "")
    return msg, stats

def update_operative(operative_id, name, email, phone, certification):
    msg, _, _ = save_operative(operative_id, name, "", email, phone, certification, "")
    return msg, msg

def delete_operative(operative_id):
    msg, tbl, _ = delete_operative_new(operative_id)
    return msg, tbl

def get_operatives_summary():
    return get_operatives_stats()

def get_operative_details(operative_id):
    _, name, nif, email, phone, cert, notes = load_operative_for_edit(operative_id)
    return f"**{name}** | {cert} | {email}"

def load_operative_for_edit_legacy(evt, current_table_data):
    try:
        row_index = evt.index[0]
        table_list = current_table_data.values.tolist() if hasattr(current_table_data, 'values') else current_table_data
        if not table_list or row_index >= len(table_list):
            return None, "", "", "", "", "*Selecione um operacional*"
        op_id = int(table_list[row_index][0])
        _, name, nif, email, phone, cert, notes = load_operative_for_edit(op_id)
        return op_id, name, email, phone, cert, f"**{name}** | {cert}"
    except Exception as e:
        return None, "", "", "", "", f"*Erro: {e}*"

# ============= FIRE ACTIONS FUNCTIONS =============
def create_fire_action(name, scheduled_date, parcel_ids, operative_ids, notes, status, photos):
    
    
    try:
        action = FireAction.objects.create(
            name=name,
            scheduled_date=scheduled_date,
            notes=notes,
            status=status if status else "Planeada"
        )
        
        if parcel_ids:
            parcels = FireParcel.objects.filter(id__in=parcel_ids)
            action.parcels.set(parcels)
            # If status is Executada, update last_burned_date on parcels
            if status == "Executada":
                parcels.update(last_burned_date=scheduled_date)
        
        if photos:
            action.notes += f"\n\n[{len(photos)} photo(s) uploaded]"
            action.save()
        
        return f"✓ Ação '{name}' criada com sucesso!", get_fire_actions_list()
    except Exception as e:
        return f"Error creating fire action: {str(e)}", get_fire_actions_list()

def get_filtered_fire_actions(date_filter, concelho_filter, estado_filter):
    """Filter actions by date, concelho and status"""
    actions = FireAction.objects.all()
    
    if date_filter and date_filter.strip():
        try:
            date_obj = datetime.strptime(date_filter.strip(), "%Y-%m-%d").date()
            actions = actions.filter(scheduled_date=date_obj)
        except ValueError:
            pass  # ignore invalid date format
    
    if concelho_filter and concelho_filter != "Todos":
        actions = actions.filter(parcels__concelho=concelho_filter).distinct()
    
    if estado_filter and estado_filter != "Todos":
        actions = actions.filter(status=estado_filter)
    
    result = []
    for action in actions:
        parcel_names = ", ".join([p.name for p in action.parcels.all()])
        result.append([
            action.id,
            action.name,
            str(action.scheduled_date),
            action.status,
            parcel_names,
        ])
    return result

def get_fire_actions_list():
    actions = FireAction.objects.all()
    result = []
    for action in actions:
        parcel_names = ", ".join([p.name for p in action.parcels.all()])
        result.append([
            action.id,
            action.name,
            str(action.scheduled_date),
            action.status,
            parcel_names,
        ])
    return result

def get_parcel_choices():
    parcels = FireParcel.objects.all()
    return [(f"{p.name} ({p.vegetation_type})", p.id) for p in parcels]

def get_operative_choices():
    operatives = Operative.objects.all()
    return [(f"{o.name} - {o.certification_level}", o.id) for o in operatives]

def get_action_by_id(action_id):
    """
    Retrieve a specific fire action by its ID
    
    Args:
        action_id: The ID of the action to retrieve
        
    Returns:
        dict: Action details or None if not found
    """
    try:
        action = FireAction.objects.get(id=action_id)
        
        return {
            'id': action.id,
            'name': action.name,
            'date': str(action.scheduled_date),
            'status': action.status,
            'execution_date': str(action.execution_date) if getattr(action, 'execution_date', None) else '',
            'parcels': [p.id for p in action.parcels.all()],
            'notes': action.notes or '',
            'photos': []  # Photos would need to be implemented separately if storing files
        }
    except FireAction.DoesNotExist:
        return None
    except Exception as e:
        print(f"Erro ao obter ação: {e}")
        return None

def update_fire_action(action_id, name, date, parcel_ids, operative_ids, notes, status, execution_date, photos):
    """
    Update an existing fire action
    """
    try:
        action = FireAction.objects.get(id=action_id)
        
        # Update basic fields
        action.name = name
        action.scheduled_date = date
        action.notes = notes
        
        old_status = action.status
        action.status = status if status else action.status
        
        # If marking as Executada, store execution date
        if status == "Executada" and execution_date:
            if hasattr(action, 'execution_date'):
                try:
                    if isinstance(execution_date, str):
                        action.execution_date = datetime.strptime(execution_date.strip(), "%Y-%m-%d").date()
                    else:
                        action.execution_date = execution_date
                except Exception:
                    pass
        
        # Update many-to-many relationships
        if parcel_ids:
            parcels = FireParcel.objects.filter(id__in=parcel_ids)
            action.parcels.set(parcels)
            # Update last_burned_date on parcels when marked as Executada
            if status == "Executada" and old_status != "Executada":
                exec_date = getattr(action, 'execution_date', None) or date
                parcels.update(last_burned_date=exec_date)
        else:
            action.parcels.clear()
        
        # Handle new photo uploads (basic implementation)
        if photos:
            action.notes += f"\n\n[{len(photos)} nova(s) foto(s) adicionada(s)]"
        
        action.save()
        return True
        
    except FireAction.DoesNotExist:
        print(f"Ação com ID {action_id} não encontrada")
        return False
    except Exception as e:
        print(f"Erro ao atualizar ação: {e}")
        return False

def delete_fire_action(action_id):
    """
    Delete a fire action from the database
    
    Args:
        action_id: ID of the action to delete
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        action = FireAction.objects.get(id=action_id)
        action_name = action.name
        action.delete()
        print(f"Ação '{action_name}' eliminada com sucesso")
        return True
        
    except FireAction.DoesNotExist:
        print(f"Ação com ID {action_id} não encontrada")
        return False
    except Exception as e:
        print(f"Erro ao eliminar ação: {e}")
        return False

# ============= PRE-PLAN FUNCTIONS =============

def get_preplan_choices():
    """Return pre-plan actions for dropdown in burning plan tab."""
    try:
        preplans = FireAction.objects.filter(status="Pre-Plano").order_by('-scheduled_date')
        return [(f"{p.name} — {p.scheduled_date} — {', '.join(pa.name for pa in p.parcels.all())}", p.id) for p in preplans]
    except Exception:
        return []

def get_preplans_list():
    try:
        actions = FireAction.objects.filter(status="Pre-Plano").order_by('scheduled_date')
        result = []
        for a in actions:
            responsible = getattr(a, 'responsible', '') or ''
            result.append([a.id, a.name, responsible, str(a.scheduled_date)])
        return result
    except Exception:
        return []

def _save_photo_files(action_id, photos):
    """Copy uploaded Gradio temp files to persistent media dir."""
    import shutil
    saved = []
    folder = MEDIA_DIR / str(action_id)
    folder.mkdir(parents=True, exist_ok=True)
    for p in (photos or []):
        try:
            src = p.name if hasattr(p, 'name') else str(p)
            ext = os.path.splitext(src)[1] or ".jpg"
            ts  = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            dst = folder / f"photo_{ts}{ext}"
            shutil.copy2(src, dst)
            saved.append(str(dst))
        except Exception as e:
            print(f"Photo save error: {e}")
    return saved

def _get_photo_paths(action_id):
    """Return sorted list of existing photo paths for a pre-plan."""
    folder = MEDIA_DIR / str(action_id)
    if not folder.exists():
        return []
    exts = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    return sorted(str(f) for f in folder.iterdir() if f.suffix.lower() in exts)

def save_preplan(preplan_id, name, responsible, indicative_date, parcel_ids, notes, photos):
    if not name or not name.strip():
        return "❌ O nome é obrigatório.", get_preplans_list()
    if not parcel_ids:
        return "❌ Selecione pelo menos uma parcela.", get_preplans_list()
    try:
        if preplan_id:
            action = FireAction.objects.get(id=int(preplan_id))
            action.name = name.strip()
            action.notes = (notes or "").strip()
            if indicative_date:
                action.scheduled_date = indicative_date
            if hasattr(action, 'responsible'):
                action.responsible = responsible or ""
            action.save()
            action.parcels.set(FireParcel.objects.filter(id__in=parcel_ids))
            if photos:
                _save_photo_files(action.id, photos)
            msg = f"✅ Pré-Plano '{name}' atualizado."
        else:
            kwargs = dict(
                name=name.strip(),
                scheduled_date=indicative_date or datetime.now().date(),
                notes=(notes or "").strip(),
                status="Pre-Plano"
            )
            if 'responsible' in [f.name for f in FireAction._meta.get_fields()]:
                kwargs['responsible'] = responsible or ""
            action = FireAction.objects.create(**kwargs)
            if parcel_ids:
                action.parcels.set(FireParcel.objects.filter(id__in=parcel_ids))
            if photos:
                _save_photo_files(action.id, photos)
            msg = f"✅ Pré-Plano '{name}' criado."
        return msg, get_preplans_list()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_preplans_list()

def load_preplan(preplan_id):
    if not preplan_id:
        return None, "", "", None, [], "", None, []
    try:
        a = FireAction.objects.get(id=int(preplan_id))
        parcel_ids = [p.id for p in a.parcels.all()]
        responsible = getattr(a, 'responsible', '') or ''
        photos = _get_photo_paths(a.id)
        return a.id, a.name, responsible, str(a.scheduled_date), parcel_ids, a.notes or "", None, photos
    except Exception:
        return None, "", "", None, [], "", None, []

def delete_preplan(preplan_id):
    if not preplan_id:
        return "❌ Selecione um pré-plano.", get_preplans_list()
    try:
        a = FireAction.objects.get(id=int(preplan_id))
        try:
            if hasattr(a, 'burningplan'):
                return f"⚠️ Não é possível eliminar — tem um Plano de Queima associado.", get_preplans_list()
        except Exception:
            pass
        name = a.name
        a.delete()
        return f"✅ Pré-Plano '{name}' eliminado.", get_preplans_list()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_preplans_list()

# ============= BURNING PLAN FUNCTIONS =============

PROBLEMS_CHOICES = [
    "Dificuldade de ignição",
    "Necessidade de rescaldo",
    "Organização deficiente",
    "Meios humanos insuficientes",
    "Equipamento insuficiente",
    "Fogo demasiado intenso",
    "Deficiente dispersão do fumo",
    "Deficiente segurança dos operadores",
    "Fuga do fogo",
    "Outros",
]

WEATHER_STATE_CHOICES = [
    "0 - Limpo (<10% nuvens)",
    "1 - Nuvens dispersas (10-50%)",
    "2 - Bastante nublado (60-90%)",
    "3 - Muito nublado (>90%)",
    "4 - Nevoeiro",
    "5 - Chuviscos",
    "6 - Chuva",
    "7 - Neve ou granizo",
    "8 - Aguaceiros",
]

WIND_DIR_CHOICES = ["N", "NE", "E", "SE", "S", "SW", "W", "NW", "V (variável)"]

FIRE_CONDUCT_CHOICES = [
    "1 - Contra o vento / contra o declive",
    "2 - Por linhas sucessivas",
    "3 - Perimetral",
    "4 - De flanco",
    "5 - Outro",
]

def get_burning_plans_list():
    try:
        plans = BurningPlan.objects.select_related("pre_plan").order_by("-execution_date")
        result = []
        for p in plans:
            parcel_names = ", ".join(pa.name for pa in p.pre_plan.parcels.all()) if p.pre_plan else ""
            result.append([p.id, p.pre_plan.name if p.pre_plan else "—",
                           str(p.execution_date), parcel_names, p.num_men or 0])
        return result
    except Exception:
        return []


def _save_bp_photo_files(bp_id, photos):
    """Save uploaded photos for a BurningPlan."""
    import shutil
    saved = []
    folder = BP_MEDIA_DIR / str(bp_id)
    folder.mkdir(parents=True, exist_ok=True)
    for p in (photos or []):
        try:
            src = p.name if hasattr(p, 'name') else str(p)
            ext = os.path.splitext(src)[1] or ".jpg"
            ts  = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
            dst = folder / f"bp_photo_{ts}{ext}"
            shutil.copy2(src, dst)
            saved.append(str(dst))
        except Exception as e:
            print(f"BP photo save error: {e}")
    return saved

def _get_bp_photo_paths(bp_id):
    """Return sorted list of photo paths for a BurningPlan."""
    if not bp_id:
        return []
    folder = BP_MEDIA_DIR / str(bp_id)
    if not folder.exists():
        return []
    exts = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
    return sorted(str(f) for f in folder.iterdir() if f.suffix.lower() in exts)

def save_burning_plan(bp_id, preplan_id, execution_date,
                      operative_ids, problems,
                      fuel_superficial, fuel_manta_f, fuel_manta_h,
                      weather_state, wind_speed_beaufort, wind_speed_kmh,
                      wind_dir, fire_conduct, fire_conduct_other,
                      num_men, vehicles_json,
                      burn_effects, burn_efficiency, extra_notes,
                      photos=None):
    if not preplan_id:
        return "❌ Selecione um Pré-Plano.", get_burning_plans_list()
    if not execution_date:
        return "❌ A data de execução é obrigatória.", get_burning_plans_list()
    try:
        pre_plan = FireAction.objects.get(id=int(preplan_id))
        problems_str = "; ".join(problems) if problems else ""
        # Extract conduct code (first char before " -")
        conduct_code = (fire_conduct or "")[:1] or ""

        data = dict(
            pre_plan=pre_plan,
            execution_date=execution_date,
            num_men=int(num_men) if num_men else None,
            problems=problems_str,
            fuel_superficial=fuel_superficial or "",
            fuel_manta_f=fuel_manta_f or "",
            fuel_manta_h=fuel_manta_h or "",
            weather_state=weather_state or "",
            wind_speed_beaufort=str(wind_speed_beaufort) if wind_speed_beaufort else "",
            wind_speed_kmh=str(wind_speed_kmh) if wind_speed_kmh else "",
            wind_direction=wind_dir or "",
            fire_conduct=conduct_code,
            fire_conduct_other=fire_conduct_other or "",
            vehicles=vehicles_json or "{}",
            burn_effects=burn_effects or "",
            burn_efficiency=burn_efficiency or "",
            notes=extra_notes or "",
        )

        if bp_id:
            bp = BurningPlan.objects.get(id=int(bp_id))
            for k, v in data.items():
                setattr(bp, k, v)
            bp.save()
            if operative_ids is not None:
                bp.operatives.set(Operative.objects.filter(id__in=operative_ids))
            if photos:
                _save_bp_photo_files(bp.id, photos)
            msg = "✅ Plano de Queima atualizado."
        else:
            bp = BurningPlan.objects.create(**data)
            if operative_ids:
                bp.operatives.set(Operative.objects.filter(id__in=operative_ids))
            if photos:
                _save_bp_photo_files(bp.id, photos)
            # Mark pre-plan as Executada
            pre_plan.status = "Executada"
            pre_plan.save()
            msg = "✅ Plano de Queima registado."

        return msg, get_burning_plans_list()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_burning_plans_list()

def load_burning_plan(bp_id):
    if not bp_id:
        return (None,) + ("",) * 18
    try:
        bp = BurningPlan.objects.get(id=int(bp_id))
        problems = bp.problems.split("; ") if bp.problems else []
        operative_ids = [o.id for o in bp.operatives.all()]
        # Rebuild full conduct string from stored code
        conduct_map = {c[:1]: c for c in FIRE_CONDUCT_CHOICES}
        fire_conduct_full = conduct_map.get(bp.fire_conduct, None) if bp.fire_conduct else None
        # Parse vehicles JSON
        import json as _json
        try:
            vd = _json.loads(bp.vehicles or "{}")
        except Exception:
            vd = {}
        return (
            bp.id, bp.pre_plan_id, str(bp.execution_date),
            operative_ids, problems,
            bp.fuel_superficial or "", bp.fuel_manta_f or "", bp.fuel_manta_h or "",
            bp.weather_state or None,
            bp.wind_speed_beaufort or "", bp.wind_speed_kmh or "",
            bp.wind_direction or None,
            fire_conduct_full, bp.fire_conduct_other or "",
            bp.num_men or "", str(vd.get("VFCI", 0)), str(vd.get("VFCM", 0)), vd.get("Outro", ""),
            bp.burn_effects or "", bp.burn_efficiency or "", bp.notes or "",
        )
    except Exception:
        return (None,) + ("",) * 20

def delete_burning_plan(bp_id):
    if not bp_id:
        return "❌ Selecione um plano.", get_burning_plans_list()
    try:
        bp = BurningPlan.objects.get(id=int(bp_id))
        # Revert pre-plan status back to Pre-Plano
        if bp.pre_plan:
            bp.pre_plan.status = "Pre-Plano"
            bp.pre_plan.save()
        bp.delete()
        return "✅ Plano de Queima eliminado.", get_burning_plans_list()
    except Exception as e:
        return f"❌ Erro: {str(e)}", get_burning_plans_list()


def generate_burning_plan_report(bp_id):
    """Generate a Word (.docx) report matching the POQ template — python-docx only."""
    import tempfile
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import date as _date

    if not bp_id:
        return None, "❌ Selecione um Plano de Queima primeiro."
    try:
        bp = BurningPlan.objects.select_related("pre_plan").get(id=int(bp_id))
    except BurningPlan.DoesNotExist:
        return None, "❌ Plano de Queima não encontrado."
    except Exception as e:
        return None, f"❌ Erro: {str(e)}"

    try:
        # ── Data collection ────────────────────────────────────────────────
        operatives      = list(bp.operatives.all())
        operative_names = [o.name for o in operatives]
        op_creds        = [getattr(o, 'certification', '') or '' for o in operatives]
        parcel_names = ""; parcel_ids_str = ""; responsible = ""
        if bp.pre_plan:
            parcels        = list(bp.pre_plan.parcels.all())
            parcel_names   = ", ".join(pa.name for pa in parcels)
            parcel_ids_str = ", ".join(str(pa.id) for pa in parcels)
            responsible    = getattr(bp.pre_plan, "responsible", "") or ""

        import json as _j
        try:
            vd = _j.loads(bp.vehicles or "{}")
        except Exception:
            vd = {}
        problems = [p.strip() for p in bp.problems.split(";") if p.strip()] if bp.problems else []
        conduct_map       = {c[:1]: c for c in FIRE_CONDUCT_CHOICES}
        fire_conduct_full = conduct_map.get(bp.fire_conduct, bp.fire_conduct or "")

        # ── Colour palette ─────────────────────────────────────────────────
        HDR_FILL = "D9D9D9"
        GREY_TXT = RGBColor(0x55, 0x55, 0x55)
        NAVY     = RGBColor(0x1F, 0x3E, 0x6B)

        def _v(val, fallback=""):
            s = str(val).strip() if val not in (None, "", "None") else ""
            return s if s else fallback

        # ── Cell helpers ───────────────────────────────────────────────────
        def _set_bg(cell, fill):
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)

        def _set_width(cell, cm):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW  = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"),    str(int(cm * 567)))
            tcW.set(qn("w:type"), "dxa")

        def _set_row_height(row, cm, exact=True):
            trPr = row._tr.get_or_add_trPr()
            trH  = OxmlElement("w:trHeight")
            trH.set(qn("w:val"),   str(int(cm * 567)))
            trH.set(qn("w:hRule"), "exact" if exact else "atLeast")
            trPr.append(trH)

        def _cell_p(cell, text, bold=False, sz=9, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            p.clear(); p.alignment = align
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(text); run.bold = bold; run.font.size = Pt(sz); run.italic = italic
            if color: run.font.color.rgb = color
            return p

        def hdr(cell, text, sz=9):
            _set_bg(cell, HDR_FILL)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_p(cell, text, bold=True, sz=sz)

        def val(cell, text, sz=9):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_p(cell, _v(text), sz=sz)

        def set_tbl_widths(tbl, widths_cm):
            """Set table total width and per-cell widths."""
            tbl_xml  = tbl._tbl
            tblPr    = tbl_xml.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr'); tbl_xml.insert(0, tblPr)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'),    str(int(sum(widths_cm) * 567)))
            tblW.set(qn('w:type'), 'dxa')
            old = tblPr.find(qn('w:tblW'))
            if old is not None: tblPr.remove(old)
            tblPr.append(tblW)
            for row in tbl.rows:
                for i, cell in enumerate(row.cells):
                    if i < len(widths_cm):
                        _set_width(cell, widths_cm[i])

        def section_heading(doc, text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(3)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single"); bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "1");      bot.set(qn("w:color"), "1F3E6B")
            pBdr.append(bot); pPr.append(pBdr)
            run = p.add_run(text); run.bold = True; run.font.size = Pt(11); run.font.color.rgb = NAVY

        def spacer(doc, pt=4):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(pt)

        # ── Document ───────────────────────────────────────────────────────
        doc = DocxDocument()
        for sec in doc.sections:
            sec.page_width   = Cm(21);   sec.page_height  = Cm(29.7)
            sec.left_margin  = Cm(2.5);  sec.right_margin  = Cm(2.5)
            sec.top_margin   = Cm(2.0);  sec.bottom_margin = Cm(2.0)
        CW = 16.0   # content width cm  (21 - 2*2.5)

        # ── Cover ──────────────────────────────────────────────────────────
        for txt, sz, sp in [("Plano Operacional de Queima", 16, 4), ("POQ - Execução", 13, 12)]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(sp)
            r = p.add_run(txt); r.bold = True; r.font.size = Pt(sz)

        tbl = doc.add_table(rows=4, cols=2); tbl.style = "Table Grid"
        set_tbl_widths(tbl, [5.5, CW - 5.5])
        hdr(tbl.rows[0].cells[0], "PLANO DE FOGO CONTROLADO N.º")
        val(tbl.rows[0].cells[1], f"{bp.id} / {_date.today().year}")
        hdr(tbl.rows[1].cells[0], "Técnico Responsável de Queima")
        val(tbl.rows[1].cells[1], responsible)
        hdr(tbl.rows[2].cells[0], "Parcela(s)")
        val(tbl.rows[2].cells[1], parcel_names)
        hdr(tbl.rows[3].cells[0], "Data de Execução")
        val(tbl.rows[3].cells[1], _v(bp.execution_date))
        spacer(doc, 8)

        # ── 1. Técnicos envolvidos ─────────────────────────────────────────
        section_heading(doc, "1. Técnicos Envolvidos na Queima")
        n_ops = max(len(operative_names), 1)
        tbl = doc.add_table(rows=1 + n_ops, cols=4); tbl.style = "Table Grid"
        cw1 = [2.0, CW * 0.38, CW * 0.28, CW - 2.0 - CW*0.38 - CW*0.28]
        set_tbl_widths(tbl, cw1)
        for txt, i in [("ID_PARC", 0), ("Nome do Técnico", 1), ("Nº Credenciação", 2), ("Funções", 3)]:
            hdr(tbl.rows[0].cells[i], txt)
        for i, name in enumerate(operative_names or [""]):
            r = tbl.rows[1 + i]
            val(r.cells[0], parcel_ids_str); val(r.cells[1], name)
            val(r.cells[2], op_creds[i] if i < len(op_creds) else ""); val(r.cells[3], "")
        spacer(doc, 6)

        # ── 2. Problemas operacionais ──────────────────────────────────────
        section_heading(doc, "2. Problemas Operacionais")
        n_prob = max(len(problems), 1)
        tbl = doc.add_table(rows=1 + n_prob, cols=3); tbl.style = "Table Grid"
        set_tbl_widths(tbl, [2.0, 3.0, CW - 5.0])
        for txt, i in [("ID_PARC", 0), ("Data", 1), ("Descrição", 2)]:
            hdr(tbl.rows[0].cells[i], txt)
        for i, prob in enumerate(problems or [""]):
            r = tbl.rows[1 + i]
            val(r.cells[0], parcel_ids_str)
            val(r.cells[1], _v(bp.execution_date))
            val(r.cells[2], prob)
        spacer(doc, 6)

        # ── 3. Humidade do combustível ─────────────────────────────────────
        section_heading(doc, "3. Humidade do Combustível")
        tbl = doc.add_table(rows=2, cols=5); tbl.style = "Table Grid"
        w5 = CW / 5
        set_tbl_widths(tbl, [w5] * 5)
        for txt, i in [("ID_PARC",0),("Data",1),("Metodologia",2),("Superficial (%)",3),("Manta morta F / H (%)",4)]:
            hdr(tbl.rows[0].cells[i], txt, sz=8)
        r = tbl.rows[1]
        val(r.cells[0], parcel_ids_str); val(r.cells[1], _v(bp.execution_date))
        val(r.cells[2], ""); val(r.cells[3], _v(bp.fuel_superficial))
        val(r.cells[4], f"{_v(bp.fuel_manta_f)} / {_v(bp.fuel_manta_h)}")
        spacer(doc, 6)

        # ── 4. Meteorologia ────────────────────────────────────────────────
        section_heading(doc, "4. Meteorologia")
        col_w = [2.0, 2.2, 2.0, 2.4, 2.0, 1.6, CW - 12.2]
        tbl = doc.add_table(rows=2, cols=7); tbl.style = "Table Grid"
        set_tbl_widths(tbl, col_w)
        for txt, i in [("ID_PARC",0),("Data",1),("Est. tempo",2),
                       ("Vel. vento\n(Bft/km/h)",3),("Dir. vento",4),("Conduç.",5),("Notas",6)]:
            hdr(tbl.rows[0].cells[i], txt, sz=8)
        r = tbl.rows[1]
        val(r.cells[0], parcel_ids_str); val(r.cells[1], _v(bp.execution_date))
        val(r.cells[2], _v(bp.weather_state))
        val(r.cells[3], f"{_v(bp.wind_speed_beaufort)} / {_v(bp.wind_speed_kmh)}")
        val(r.cells[4], _v(bp.wind_direction))
        val(r.cells[5], _v(bp.fire_conduct))
        val(r.cells[6], _v(bp.fire_conduct_other))
        spacer(doc, 6)

        # ── 5. Esquema de condução da queima ───────────────────────────────
        # Exact structure from esquema_POQ.odt:
        # 6 underlying cols: A=5.166, B=1.498, C=0.813, D=2.506, E=2.579, F=2.819  total=15.381cm
        # Row 0: [A+B=6.664 ID_PARC] [C+D=3.319 Início] [E+F=5.398 Fim]
        # Row 1: [A=5.166 Meios utilizados] [B+C=2.311 Téc.n.º] [D=2.506 VLCI(<500l)] [E=2.579 Homens:] [F=2.819 VFCI(_l)]
        # Row 2: [all 6 cols = 15.381, height 9cm, ESQUEMA: label top-left + instruction, rest empty]
        section_heading(doc, "5. Esquema de Condução da Queima")

        COL_A, COL_B, COL_C = 5.166, 1.498, 0.813
        COL_D, COL_E, COL_F = 2.506, 2.579, 2.819
        SCH_TOTAL = COL_A + COL_B + COL_C + COL_D + COL_E + COL_F  # 15.381

        # Build a 6-column table
        tbl = doc.add_table(rows=3, cols=6)
        tbl.style = "Table Grid"
        set_tbl_widths(tbl, [COL_A, COL_B, COL_C, COL_D, COL_E, COL_F])

        # ── Row 0: ID_PARC | Início | Fim ────────────────────────────────
        # Merge A+B for ID_PARC
        tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])
        hdr(tbl.rows[0].cells[0], f"ID_PARC: {parcel_ids_str}")
        # Merge C+D for Início
        tbl.rows[0].cells[2].merge(tbl.rows[0].cells[3])
        hdr(tbl.rows[0].cells[2], "Início:       horas e       min")
        # Merge E+F for Fim
        tbl.rows[0].cells[4].merge(tbl.rows[0].cells[5])
        hdr(tbl.rows[0].cells[4], "Fim:       horas e       min")

        # ── Row 1: Meios utilizados | Téc.n.º | VLCI | Homens | VFCI ────
        # Col A alone: Meios utilizados
        hdr(tbl.rows[1].cells[0], "Meios utilizados")
        # Merge B+C for Téc.n.º value
        tbl.rows[1].cells[1].merge(tbl.rows[1].cells[2])
        tec_val = _v(str(len(operative_names)), "___")
        val(tbl.rows[1].cells[1], f"Téc.n.º: {tec_val}")
        # Col D: VLCI (<500l)
        vlci_val = _v(str(vd.get("VLCI", vd.get("VFCM", 0))), "___")
        val(tbl.rows[1].cells[3], f"VLCI: {vlci_val}\n(<500l)")
        # Col E: Homens
        hom_val = _v(str(bp.num_men or ""), "___")
        val(tbl.rows[1].cells[4], f"Homens: {hom_val}")
        # Col F: VFCI
        vfci_val = _v(str(vd.get("VFCI", 0)), "___")
        val(tbl.rows[1].cells[5], f"VFCI:\n({vfci_val}____l)")

        # ── Row 2: full-width tall ESQUEMA drawing cell ────────────────────
        # Merge all 6 columns
        for ci in range(1, 6):
            tbl.rows[2].cells[0].merge(tbl.rows[2].cells[ci])
        schema_cell = tbl.rows[2].cells[0]

        # Set exact height of 9 cm (matching ODT row-height=9.001cm)
        trPr = tbl.rows[2]._tr.get_or_add_trPr()
        trH  = OxmlElement("w:trHeight")
        trH.set(qn("w:val"),   str(int(9 * 567)))   # 9cm in twips
        trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)

        # ESQUEMA label: bold "ESQUEMA" + colon + small instruction text (top-left, like ODT)
        schema_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        sp = schema_cell.paragraphs[0]
        sp.clear()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after  = Pt(0)
        # Bold "ESQUEMA"
        r_bold = sp.add_run("ESQUEMA")
        r_bold.bold = True; r_bold.font.size = Pt(9)
        # Colon + instruction in normal weight small text
        r_inst = sp.add_run(": Indicar os sentidos do declive, Norte, direcção do fogo, "                             "a localização da ignição (A), das faixas de segurança (B), "                             "localização da foto tirada (C)...")
        r_inst.font.size = Pt(8)

        # Legend below table
        leg_p = doc.add_paragraph()
        leg_p.paragraph_format.space_before = Pt(2)
        leg_p.paragraph_format.space_after  = Pt(2)
        r_leg = leg_p.add_run("VLCI – Veículo Ligeiro de Combate a Incêndio          "                               "VFCI – Veículo Florestal de Combate a Incêndio")
        r_leg.font.size = Pt(8); r_leg.italic = True

        spacer(doc, 4)

        # ── 6. Efeitos imediatos da queima ─────────────────────────────────
        section_heading(doc, "6. Efeitos Imediatos da Queima")
        tbl = doc.add_table(rows=2, cols=6); tbl.style = "Table Grid"
        ew = CW / 6
        set_tbl_widths(tbl, [ew] * 6)
        for txt, i in [("ID_PARC",0),("Vegetação\n(área %)",1),("Manta morta\n(área %)",2),
                       ("% Copas\nconsumidas",3),("Ø calcinado\n(mm)",4),("Redução\nespessura %",5)]:
            hdr(tbl.rows[0].cells[i], txt, sz=8)
        r = tbl.rows[1]
        val(r.cells[0], parcel_ids_str)
        for i in range(1, 6):
            val(r.cells[i], "")
        spacer(doc, 6)

        # ── 7. Eficácia da queima ──────────────────────────────────────────
        section_heading(doc, "7. Eficácia da Queima")
        tbl = doc.add_table(rows=2, cols=5); tbl.style = "Table Grid"
        fw = CW / 5
        set_tbl_widths(tbl, [fw] * 5)
        for txt, i in [("ID_PARC",0),("Resp. objectivos",1),("Redução combustível",2),
                       ("Cond. meteorológicas",3),("Porquê?",4)]:
            hdr(tbl.rows[0].cells[i], txt, sz=8)
        r = tbl.rows[1]
        val(r.cells[0], parcel_ids_str)
        val(r.cells[1], _v(bp.burn_efficiency))
        val(r.cells[2], "")
        val(r.cells[3], _v(bp.weather_state))
        val(r.cells[4], _v(bp.burn_effects))
        spacer(doc, 6)

        # ── 8. Notas adicionais ────────────────────────────────────────────
        if bp.notes and bp.notes.strip():
            section_heading(doc, "8. Notas Adicionais")
            tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
            set_tbl_widths(tbl, [CW])
            val(tbl.rows[0].cells[0], bp.notes)
            spacer(doc, 6)

        # ── 9. Índice de fotografias ───────────────────────────────────────
        section_heading(doc, "9. Índice de Fotografias")
        photo_paths = _get_bp_photo_paths(bp.id)
        if photo_paths:
            photos_to_show = photo_paths[:6]
            rows_needed    = (len(photos_to_show) + 1) // 2
            tbl = doc.add_table(rows=rows_needed, cols=2); tbl.style = "Table Grid"
            set_tbl_widths(tbl, [CW / 2, CW / 2])
            for idx_, ph_path in enumerate(photos_to_show):
                ph_cell = tbl.rows[idx_ // 2].cells[idx_ % 2]
                ph_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                ph_p = ph_cell.paragraphs[0]; ph_p.clear()
                ph_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    run_ = ph_p.add_run()
                    run_.add_picture(ph_path, width=Inches(2.8))
                except Exception:
                    ph_p.add_run(f"[Foto {idx_+1}: erro ao carregar]").font.size = Pt(8)
                cap = ph_cell.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(f"Foto {idx_+1}"); cr.font.size = Pt(8); cr.italic = True
        else:
            tbl = doc.add_table(rows=2, cols=4); tbl.style = "Table Grid"
            set_tbl_widths(tbl, [2.0, 2.0, CW * 0.3, CW - 4.0 - CW * 0.3])
            for txt, i in [("ID_PARC",0),("Data",1),("Foto — Antes da queima",2),("Foto — Depois da queima",3)]:
                hdr(tbl.rows[0].cells[i], txt, sz=8)
            for i in range(4):
                val(tbl.rows[1].cells[i], "")
        spacer(doc, 6)

        # ── 10. Assinaturas ────────────────────────────────────────────────
        section_heading(doc, "10. Assinaturas")
        tbl = doc.add_table(rows=2, cols=2); tbl.style = "Table Grid"
        set_tbl_widths(tbl, [CW / 2, CW / 2])
        hdr(tbl.rows[0].cells[0], "Técnico Responsável de Queima")
        hdr(tbl.rows[0].cells[1], "Técnico de Fogo Controlado")
        for c in tbl.rows[1].cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_ = c.paragraphs[0]; p_.clear()
            p_.paragraph_format.space_before = Pt(28); p_.paragraph_format.space_after = Pt(28)
            r_ = p_.add_run("Assinatura: _______________________________     Data: ___/___/______")
            r_.font.size = Pt(9); r_.font.color.rgb = GREY_TXT

        # ── Save ───────────────────────────────────────────────────────────
        tmp_dir  = Path(tempfile.mkdtemp())
        out_path = tmp_dir / f"POQ_{bp.id}_{_date.today().strftime('%Y%m%d')}.docx"
        doc.save(str(out_path))
        return str(out_path), f"✅ Relatório POQ gerado: POQ_{bp.id}.docx"

    except Exception as e:
        import traceback
        return None, f"❌ Erro ao gerar relatório: {str(e)}\n{traceback.format_exc()}"

# ============= CALENDAR & WEATHER FUNCTIONS =============
def generate_calendar_view(year, month):
    actions = FireAction.objects.filter(
        scheduled_date__year=year,
        scheduled_date__month=month
    )
    
    action_dates = {}
    for action in actions:
        date_key = action.scheduled_date.day
        if date_key not in action_dates:
            action_dates[date_key] = []
        action_dates[date_key].append(action.name)
    
    cal_obj = cal.monthcalendar(year, month)
    month_name = cal.month_name[month]
    
    html = f"<h3>{month_name} {year}</h3><table style='width:100%; border-collapse: collapse;'>"
    html += "<tr><th>Mon</th><th>Tue</th><th>Wed</th><th>Thu</th><th>Fri</th><th>Sat</th><th>Sun</th></tr>"
    
    for week in cal_obj:
        html += "<tr>"
        for day in week:
            if day == 0:
                html += "<td style='padding:10px; border:1px solid #ddd;'></td>"
            else:
                style = "padding:10px; border:1px solid #ddd;"
                if day in action_dates:
                    style += "background-color:#000000;"
                html += f"<td style='{style}'><strong>{day}</strong>"
                if day in action_dates:
                    for action_name in action_dates[day]:
                        html += f"<br/><small>🔥 {action_name}</small>"
                html += "</td>"
        html += "</tr>"
    html += "</table>"
    
    return html

def get_weather_forecast(location="Algarve, PT"):
    try:
        lat, lon = 37.2555, -8.3965
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&daily=temperature_2m_max,temperature_2m_min,precipitation_sum,windspeed_10m_max&timezone=auto&forecast_days=7"
        response = requests.get(url, timeout=5)
        data = response.json()
        
        daily = data['daily']
        forecast_html = "<h3>7-Day Weather Forecast</h3>"
        forecast_html += "<table style='width:100%; border-collapse: collapse;'>"
        forecast_html += "<tr><th>Date</th><th>Temp (°C)</th><th>Precipitation (mm)</th><th>Wind (km/h)</th></tr>"
        
        for i in range(len(daily['time'])):
            date = daily['time'][i]
            temp_max = daily['temperature_2m_max'][i]
            temp_min = daily['temperature_2m_min'][i]
            precip = daily['precipitation_sum'][i]
            wind = daily['windspeed_10m_max'][i]
            
            forecast_html += f"""
            <tr style='border:1px solid #ddd;'>
                <td style='padding:8px;'>{date}</td>
                <td style='padding:8px;'>{temp_min}° - {temp_max}°</td>
                <td style='padding:8px;'>{precip}</td>
                <td style='padding:8px;'>{wind}</td>
            </tr>
            """
        forecast_html += "</table>"
        return forecast_html
    except Exception as e:
        return f"<p>Unable to fetch weather data: {str(e)}</p>"

def get_hourly_weather(location_json):
    if not location_json:
        return pd.DataFrame()

    loc = json.loads(location_json)
    lat, lon = loc["lat"], loc["lon"]

    url = (
        "https://api.open-meteo.com/v1/forecast"
        f"?latitude={lat}&longitude={lon}"
        "&hourly=temperature_2m,relative_humidity_2m,"
        "precipitation,wind_speed_10m,wind_gusts_10m"
        "&forecast_days=1"
        "&timezone=Europe%2FLisbon"
    )

    r = requests.get(url, timeout=5)
    data = r.json()["hourly"]

    df = pd.DataFrame({
        "Hora": pd.to_datetime(data["time"]).dt.strftime("%H:%M"),
        "Temp (°C)": data["temperature_2m"],
        "HR (%)": data["relative_humidity_2m"],
        "Chuva (mm)": data["precipitation"],
        "Vento (km/h)": data["wind_speed_10m"],
        "Rajadas (km/h)": data["wind_gusts_10m"],
    })

    return df

def load_hourly_weather(location_json):
    if not location_json:
        return (
            "Clique numa cidade no mapa",
            pd.DataFrame()
        )

    loc = json.loads(location_json)

    url = (
        "https://api.open-meteo.com/v1/forecast"
        f"?latitude={loc['lat']}&longitude={loc['lon']}"
        "&hourly=temperature_2m,relative_humidity_2m,"
        "precipitation,wind_speed_10m,wind_gusts_10m"
        "&forecast_days=1"
        "&timezone=Europe%2FLisbon"
    )

    r = requests.get(url, timeout=5)
    h = r.json()["hourly"]

    df = pd.DataFrame({
        "Hora": pd.to_datetime(h["time"]).strftime("%H:%M"),
        "Temp (°C)": h["temperature_2m"],
        "HR (%)": h["relative_humidity_2m"],
        "Chuva (mm)": h["precipitation"],
        "Vento (km/h)": h["wind_speed_10m"],
        "Rajadas (km/h)": h["wind_gusts_10m"],
    })

    return f"📍 {loc['name']}", df

# ============= REPORTS FUNCTIONS =============
def generate_report(report_type):
    if report_type == "Resumo Parcelas":
        parcels = FireParcel.objects.all()
        data = {
            'Name': [p.name for p in parcels],
            'Vegetation': [p.vegetation_type for p in parcels],
            'Risk Level': [p.risk_level for p in parcels],
        }
        df = pd.DataFrame(data)
        summary = f"**Total Parcels:** {len(parcels)}\n\n**Risk Distribution:**\n"
        risk_counts = parcels.values('risk_level').annotate(count=Count('risk_level'))
        for item in risk_counts:
            summary += f"- {item['risk_level']}: {item['count']}\n"
        return summary, df
    
    elif report_type == "Resumo Operacionais":
        operatives = Operative.objects.all()
        data = {
            'Name': [o.name for o in operatives],
            'Email': [o.email for o in operatives],
            'Certification': [o.certification_level for o in operatives],
            'Planos de Queima': [o.burning_plans.count() for o in operatives],
        }
        df = pd.DataFrame(data)
        summary = f"**Total Operatives:** {len(operatives)}\n\n**Certification Distribution:**\n"
        cert_counts = operatives.values('certification_level').annotate(count=Count('certification_level'))
        for item in cert_counts:
            summary += f"- {item['certification_level']}: {item['count']}\n"
        return summary, df
    
    elif report_type == "Resumo Ação Fogo":
        actions = FireAction.objects.all()
        data = {
            'Name': [a.name for a in actions],
            'Date': [str(a.scheduled_date) for a in actions],
            'Status': [a.status for a in actions],
            'Parcels': [a.parcels.count() for a in actions],
            'Operatives': [a.operatives.count() for a in actions],
        }
        df = pd.DataFrame(data)
        summary = f"**Total Fire Actions:** {len(actions)}\n\n**Status Distribution:**\n"
        status_counts = actions.values('status').annotate(count=Count('status'))
        for item in status_counts:
            summary += f"- {item['status']}: {item['count']}\n"
        return summary, df
    
    return "Select a report type", pd.DataFrame()

def export_data(export_type, file_format):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if export_type == "Parcelas":
        parcels = FireParcel.objects.all()
        data = {'ID': [p.id for p in parcels], 'Name': [p.name for p in parcels], 'Vegetation': [p.vegetation_type for p in parcels], 'Risk Level': [p.risk_level for p in parcels]}
        df = pd.DataFrame(data)
        filename = f"parcels_export_{timestamp}.csv"
    elif export_type == "Operacionais":
        operatives = Operative.objects.all()
        data = {'ID': [o.id for o in operatives], 'Name': [o.name for o in operatives], 'Email': [o.email for o in operatives], 'Certification': [o.certification_level for o in operatives]}
        df = pd.DataFrame(data)
        filename = f"operatives_export_{timestamp}.csv"
    elif export_type == "Ações Fogo":
        actions = FireAction.objects.all()
        data = {'ID': [a.id for a in actions], 'Name': [a.name for a in actions], 'Date': [str(a.scheduled_date) for a in actions], 'Status': [a.status for a in actions]}
        df = pd.DataFrame(data)
        filename = f"fire_actions_export_{timestamp}.csv"
    else:
        return None
    
    filepath = f"/tmp/{filename}"
    if file_format == "CSV":
        df.to_csv(filepath, index=False)
    else:
        filepath = filepath.replace('.csv', '.json')
        df.to_json(filepath, orient='records', indent=2)
    return filepath

# ============= GRADIO APP WITH AUTHENTICATION =============
HEAD_JS = """
<style>
    /* Target only the Tab navigation buttons */
    .tabs .tab-nav button.tabbutton {
        color: #d35400 !important;   /* Orange text */
        font-weight: bold !important; /* Bold text */
        font-size: 1.2em !important;  /* Larger size */
        text-transform: uppercase;    /* Uppercase */
        border-bottom: 2px solid transparent;
    }

    /* Style for the currently active/selected tab */
    .tabs .tab-nav button.tabbutton.selected {
        color: #d35400 !important;
        border-bottom: 3px solid #d35400 !important; /* Orange underline for active tab */
        background-color: rgba(211, 84, 0, 0.1) !important; /* Light orange background */
    }
    
    /* Hover effect */
    .tabs .tab-nav button.tabbutton:hover {
        background-color: rgba(211, 84, 0, 0.05) !important;
    }

    .hidden-textbox { display: none !important; }
</style>
<script>
let parcelsLoadAttempts = 0;
const maxAttempts = 10;

function sendParcelsToMap() {
    // Support multiple map iframes
    const registryIframe = document.getElementById('registry-map-iframe');
    const editIframe = document.getElementById('edit-map-iframe');
    const oldIframe = document.getElementById('map-iframe'); // Keep for backward compatibility
    
    const container = document.getElementById('geojson-data-input');
    const textarea = container ? container.querySelector('textarea') : null;

    console.log('Attempt', parcelsLoadAttempts, '- registry:', !!registryIframe, 'edit:', !!editIframe, 'old:', !!oldIframe, 'textarea:', !!textarea);

    if (textarea && textarea.value) {
        try {
            const geojson = JSON.parse(textarea.value);
            console.log('Sending parcels to maps:', geojson);
            
            // Send to all available map iframes
            if (registryIframe && registryIframe.contentWindow) {
                registryIframe.contentWindow.postMessage({
                    type: 'load_parcels',
                    geojson: geojson
                }, '*');
                console.log('Parcels sent to registry map');
            }
            
            if (editIframe && editIframe.contentWindow) {
                editIframe.contentWindow.postMessage({
                    type: 'load_parcels',
                    geojson: geojson
                }, '*');
                console.log('Parcels sent to edit map');
            }
            
            // Keep old map support
            if (oldIframe && oldIframe.contentWindow) {
                oldIframe.contentWindow.postMessage({
                    type: 'load_parcels',
                    geojson: geojson
                }, '*');
                console.log('Parcels sent to old map');
            }
            
        } catch (e) {
            console.error('Error parsing or sending GeoJSON:', e);
            if (parcelsLoadAttempts < maxAttempts) {
                parcelsLoadAttempts++;
                setTimeout(sendParcelsToMap, 1000);
            }
        }
    } else {
        if (parcelsLoadAttempts < maxAttempts) {
            parcelsLoadAttempts++;
            setTimeout(sendParcelsToMap, 1000);
        } else {
            console.error('Failed to load parcels after', maxAttempts, 'attempts');
        }
    }
}

// Listen for messages from the iframes
window.addEventListener("message", (event) => {
    console.log('Received message:', event.data);
    
    // When any iframe requests parcels
    if (event.data && event.data.type === "request_parcels") {
        console.log('Iframe requested parcels');
        parcelsLoadAttempts = 0;
        sendParcelsToMap();
    }
    
    // When user draws a parcel on the map
    if (event.data && event.data.type === "parcel_geometry") {
        console.log('Received parcel geometry from map');
        
        // Find the appropriate geometry textarea
        // First try to find by elem_id
        let geometryInput = document.getElementById('geometry-input');
        let editGeometryInput = document.getElementById('geometry-input-edit');
        
        // Determine which textarea to update based on which is visible
        let targetTextarea = null;
        
        if (geometryInput) {
            const textarea = geometryInput.querySelector('textarea');
            if (textarea && textarea.offsetParent !== null) { // visible
                targetTextarea = textarea;
            }
        }
        
        if (!targetTextarea && editGeometryInput) {
            const textarea = editGeometryInput.querySelector('textarea');
            if (textarea && textarea.offsetParent !== null) { // visible
                targetTextarea = textarea;
            }
        }
        
        // Fallback: find any visible textarea with "Geometria" or "Geometry" in nearby text
        if (!targetTextarea) {
            const textareas = document.querySelectorAll('textarea');
            for (let ta of textareas) {
                if (ta.offsetParent !== null) { // is visible
                    const block = ta.closest('.block');
                    if (block && (block.textContent.includes('Geometria') || block.textContent.includes('Geometry'))) {
                        targetTextarea = ta;
                        break;
                    }
                }
            }
        }
        
        if (targetTextarea) {
            const nativeInputValueSetter = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, 'value').set;
            nativeInputValueSetter.call(targetTextarea, JSON.stringify(event.data.data));
            targetTextarea.dispatchEvent(new Event('input', { bubbles: true }));
            console.log('Geometry set in textarea');

            // Also populate area_ha and concelho if provided
            function setGradioTextbox(elemId, value) {
                if (value === null || value === undefined) return;
                const el = document.getElementById(elemId);
                if (!el) return;
                const ta = el.querySelector('textarea') || el.querySelector('input');
                if (!ta) return;
                const setter = Object.getOwnPropertyDescriptor(
                    ta.tagName === 'TEXTAREA' ? window.HTMLTextAreaElement.prototype : window.HTMLInputElement.prototype,
                    'value'
                ).set;
                setter.call(ta, String(value));
                ta.dispatchEvent(new Event('input', { bubbles: true }));
            }
            // Registry tab
            if (event.data.area_ha !== undefined) setGradioTextbox('area-ha-input', event.data.area_ha);
            if (event.data.concelho) setGradioTextbox('concelho-input', event.data.concelho);
            // Edit tab
            if (event.data.area_ha !== undefined) setGradioTextbox('edit-area-ha-input', event.data.area_ha);
            if (event.data.concelho) setGradioTextbox('edit-concelho-input', event.data.concelho);
        } else {
            console.warn('Could not find geometry textarea');
        }
    }
});

// Start attempting to load parcels when page loads
document.addEventListener('DOMContentLoaded', function() {
    console.log('DOM loaded, waiting for components...');
    setTimeout(() => {
        console.log('Starting initial parcel load');
        sendParcelsToMap();
    }, 2000);
});

// Also try when window loads
window.addEventListener('load', function() {
    console.log('Window loaded');
    setTimeout(sendParcelsToMap, 2000);
});

// Handle zoom to geometry for edit mode
window.addEventListener("message", (event) => {
    if (event.data.type === 'edit_geometry_selected') {
        const editIframe = document.getElementById('edit-map-iframe');
        const oldIframe = document.getElementById('map-iframe');
        
        const iframe = editIframe || oldIframe;
        if (iframe && iframe.contentWindow) {
            iframe.contentWindow.postMessage({
                type: 'zoom_to_geometry',
                geometry: JSON.parse(event.data.data)
            }, '*');
        }
    }
});

function triggerZoom() {
    const geomInput = document.getElementById('geometry-input-edit');
    if (geomInput) {
        const textarea = geomInput.querySelector('textarea');
        if (textarea && textarea.value) {
            try {
                const geometry = JSON.parse(textarea.value);
                const editIframe = document.getElementById('edit-map-iframe');
                if (editIframe && editIframe.contentWindow) {
                    editIframe.contentWindow.postMessage({
                        type: 'zoom_to_geometry',
                        geometry: geometry
                    }, '*');
                }
            } catch (e) {
                console.error('Error parsing geometry for zoom:', e);
            }
        }
    }
}

// Function to handle registry map actions (save/clear)
function handleRegistryMapAction(action) {
    console.log('Registry map action:', action);
    const iframe = document.getElementById('registry-map-iframe');
    if (!iframe || !iframe.contentWindow) {
        console.error('Registry map iframe not found');
        return;
    }
    
    if (action === 'save') {
        // Clear drawn items
        iframe.contentWindow.postMessage({
            type: 'clear_drawn_items'
        }, '*');
        
        // Reset view to show all parcels
        setTimeout(() => {
            iframe.contentWindow.postMessage({
                type: 'reset_view'
            }, '*');
        }, 150);
    } else if (action === 'clear') {
        // Just clear drawn items
        iframe.contentWindow.postMessage({
            type: 'clear_drawn_items'
        }, '*');
    }
}

window.addEventListener("message", (event) => {
    if (event.data?.type === "weather_location_selected") {
        const payload = event.data.payload;

        const input = document.getElementById("weather-location-input")
            ?.querySelector("textarea");

        if (input) {
            input.value = JSON.stringify(payload);
            input.dispatchEvent(new Event("input", { bubbles: true }));
        }
    }
});

window.addEventListener("message", (event) => {
    console.log("Message received:", event.data);
    if (event.data?.type === "hourly_weather_request") {
        const el = document
            .getElementById("hourly-location")
            ?.querySelector("textarea");

        if (el) {
            el.value = JSON.stringify(event.data.payload);
            el.dispatchEvent(new Event("input", { bubbles: true }));
        }
    }
});

</script>
"""

def create_main_interface():
    """Interface principal (requer autenticação)"""
    with gr.Blocks() as main_app:
        geojson_data_hidden = gr.Textbox(
            visible=True,
            elem_id="geojson-data-input", 
            elem_classes="hidden-textbox"
        )

                
        main_app.load(get_parcels_geojson, outputs=[geojson_data_hidden])
        
        # Header com info do utilizador e logout
        gr.Image(
            value=str(Path(BASE_DIR) / "images" / "banner_FBAlgarve.png"), 
            show_label=False, 
            container=False, 
            interactive=False,
            elem_id="top-banner"
        )

        gr.Markdown("---")
        with gr.Tabs():
            with gr.Tab("INICIO"):
                gr.Markdown("## Bem vindo ao Sistema de Gestão de Fogo Controlado")
                with gr.Row():
                    gr.Markdown(f"**PARCELAS:** {FireParcel.objects.count()}")
                    gr.Markdown(f"**OPERACIONAIS:** {Operative.objects.count()}")
                    gr.Markdown(f"**AÇÕES:** {FireAction.objects.count()}")
                
                with gr.Row():
                    with gr.Column():
                        gr.Markdown("""
                    O fogo controlado é uma das ferramentas mais eficazes na prevenção de incêndios rurais. Ao reduzir a acumulação de combustível vegetal, contribui para diminuir o risco de ignição e a intensidade dos fogos silvestres, protegendo pessoas, habitações e ecossistemas de elevado valor natural.

                    No Algarve, região particularmente vulnerável ao fogo devido ao clima mediterrânico e à predominância de espécies de elevada inflamabilidade, o fogo controlado assume uma relevância estratégica como instrumento de ordenamento da paisagem e de reforço da resiliência das comunidades.

                    A plataforma **Fogo Bom Algarve** centraliza o registo georreferenciado das parcelas de intervenção, o planeamento das ações de fogo controlado e a gestão das equipas certificadas, assegurando a rastreabilidade integral das operações realizadas.
                    """)
                    with gr.Column():
                        now = datetime.now()
                        calendar_display = gr.HTML(generate_calendar_view(now.year, now.month))
                        with gr.Row():
                            cal_month = gr.Dropdown(choices=[(cal.month_name[i], i) for i in range(1, 13)], value=now.month, label="Mês")
                            cal_year = gr.Number(value=now.year, label="Ano", precision=0)
                        with gr.Row():
                            update_cal_btn = gr.Button("Atualizar Calendário")
                
                
                update_cal_btn.click(generate_calendar_view, inputs=[cal_year, cal_month], outputs=calendar_display)
            
            with gr.Tab("PARCELAS-Registo"):
                gr.Markdown("# Registo de novas parcelas")
                
                with gr.Row():
                    with gr.Column(scale=2):
                        gr.HTML(REGISTRY_MAP_IFRAME)
                        geometry = gr.Textbox(label="Geometria JSON", elem_id="geometry-input", lines=2)

                    with gr.Column(scale=1):
                        p_name = gr.Textbox(label="Designação")
                        p_resp = gr.Textbox(label="Responsável")
                        p_email = gr.Textbox(label="Email")
                        p_veg = gr.Dropdown(choices=["Matos", "Pinhal", "Eucaliptal", "Misto"], label="Vegetação")
                        p_infra = gr.Dropdown(choices=["Sem informação","Inexistente", "Operacional", "Manutenção necessária"], label="Infraestrutura")
                        p_owner = gr.Radio(choices=["Sim", "Não"], label="Info Proprietário")
                        p_concelho = gr.Textbox(label="Concelho — obtido automaticamente", elem_id="concelho-input", interactive=False, elem_classes="autofill-box")
                        p_area = gr.Textbox(label="Área (ha) — calculada automaticamente", elem_id="area-ha-input", interactive=False, elem_classes="autofill-box")
                
                with gr.Row():
                        save_btn = gr.Button("💾 Guardar Parcela", variant="primary")
                        clear_btn = gr.Button("🗑️ Limpar Dados")

                output = gr.Textbox(label="Estado", interactive=False)
                
                # Hidden state for map action
                registry_action = gr.Textbox(elem_id="registry-action", visible=False)

                # Save button
                save_btn.click(
                    save_parcel,
                    inputs=[p_name, p_resp, p_email, p_veg, p_infra, p_owner, p_concelho, geometry, p_area],
                    outputs=[output, geojson_data_hidden, registry_action, p_name, p_resp, p_email, p_veg, p_infra, p_owner, p_concelho, geometry, p_area],
                    js="""
                    (name, resp, email, veg, infra, owner, concelho, geom, area) => {
                        setTimeout(() => handleRegistryMapAction('save'), 300);
                        const geomEl = document.getElementById('geometry-input');
                        const domGeom = geomEl ? (geomEl.querySelector('textarea') || {}).value : null;
                        const finalGeom = (domGeom && domGeom.trim()) ? domGeom : geom;
                        const areaEl = document.getElementById('area-ha-input');
                        const domArea = areaEl ? (areaEl.querySelector('textarea') || areaEl.querySelector('input') || {}).value : null;
                        const finalArea = (domArea && domArea.trim()) ? domArea : area;
                        const concelhoEl = document.getElementById('concelho-input');
                        const domConcelho = concelhoEl ? (concelhoEl.querySelector('textarea') || concelhoEl.querySelector('input') || {}).value : null;
                        const finalConcelho = (domConcelho && domConcelho.trim()) ? domConcelho : concelho;
                        return [name, resp, email, veg, infra, owner, finalConcelho, finalGeom, finalArea];
                    }
                    """
                )

                # Clear button
                clear_btn.click(
                    clear_registry_form,
                    inputs=None,
                    outputs=[p_name, p_resp, p_email, p_veg, p_infra, p_owner, p_concelho, geometry, output, geojson_data_hidden, registry_action, p_area],
                    js="""
                    () => {
                        setTimeout(() => handleRegistryMapAction('clear'), 100);
                        return [];
                    }
                    """
                )

                
            with gr.Tab("PARCELAS-Editar"):
                gr.Markdown("## Edição de parcelas registadas")
    
                with gr.Row():
                    filter_concelho = gr.Dropdown(
                        choices=["Todos"] + ALGARVE_CONCELHOS,
                        value="Todos",
                        label="Filtrar por Concelho",
                        scale=2
                    )
                    filter_resp = gr.Textbox(label="Filtrar por Responsável", scale=2, placeholder="Nome...")
                    filter_veg = gr.Dropdown(
                        choices=["Todos", "Matos", "Pinhal", "Eucaliptal", "Misto"],
                        value="Todos",
                        label="Filtrar por Vegetação",
                        scale=2
                    )
                    filter_parcel_btn = gr.Button("🔍 Filtrar", scale=1)

                with gr.Row():
                    with gr.Column(scale=2):
                        parcel_dropdown = gr.Dropdown(
                            choices=[(f"{p[0]} - {p[1]}", p[0]) for p in get_parcels_list()],
                            label="Selecione uma parcela para editar",
                            interactive=True,
                            value=None
                        )
                        gr.HTML(EDIT_MAP_IFRAME)
                        edit_geom = gr.Textbox(
                            label="Geometria JSON", 
                            elem_id="geometry-input-edit", 
                            lines=2,
                            placeholder="A geometria será carregada automaticamente..."
                        )

                    with gr.Column(scale=1):
                        edit_id = gr.Number(label="ID Parcela", visible=False, interactive=False)
                        e_name = gr.Textbox(label="Designação *", interactive=True)
                        e_resp = gr.Textbox(label="Responsável *", interactive=True)
                        e_email = gr.Textbox(label="Email", interactive=True)
                        e_veg = gr.Dropdown(
                            choices=["Matos", "Pinhal", "Eucaliptal", "Misto"], 
                            label="Vegetação",
                            interactive=True
                        )
                        e_infra = gr.Dropdown(
                            choices=["Sem informação", "Inexistente", "Operacional", "Manutenção necessária"], 
                            label="Infraestrutura",
                            interactive=True
                        )
                        e_owner = gr.Radio(
                            choices=["Sim", "Não"],
                            label="Info Proprietário",
                            interactive=True
                        )
                        e_concelho = gr.Textbox(label="Concelho — obtido automaticamente", elem_id="edit-concelho-input", interactive=False, elem_classes="autofill-box")
                        e_area = gr.Textbox(label="Área (ha) — calculada automaticamente", elem_id="edit-area-ha-input", interactive=False, elem_classes="autofill-box")
                with gr.Row():
                    btn_update = gr.Button("💾 Atualizar", variant="primary")
                    btn_delete = gr.Button("🗑️ Eliminar", variant="stop")
    
                edit_msg = gr.Markdown()
                zoom_trigger = gr.Textbox(visible=False, elem_id="zoom-geom-trigger")  # holds geometry JSON for zoom
                edit_map_loader = gr.HTML()  # For map reload scripts

                # Event handlers
                def apply_parcel_filters(concelho, resp, veg):
                    filtered = get_filtered_parcels(concelho, resp, veg)
                    return gr.update(choices=filtered, value=None)

                filter_parcel_btn.click(
                    fn=apply_parcel_filters,
                    inputs=[filter_concelho, filter_resp, filter_veg],
                    outputs=[parcel_dropdown]
                )

                parcel_dropdown.change(
                    fn=load_parcel_for_edit_dropdown,
                    inputs=[parcel_dropdown],
                    outputs=[edit_id, e_name, e_resp, e_email, e_veg, e_infra, e_owner, e_concelho, e_area, edit_geom, edit_msg]
                ).then(
                    fn=zoom_to_parcel,
                    inputs=[parcel_dropdown],
                    outputs=[zoom_trigger]
                ).then(
                    fn=None,
                    inputs=[zoom_trigger],
                    outputs=None,
                    js="""
                    (geomJson) => {
                        if (!geomJson) return;
                        try {
                            const geometry = JSON.parse(geomJson);
                            function trySend(attempts) {
                                const iframe = document.getElementById('edit-map-iframe');
                                if (iframe && iframe.contentWindow) {
                                    iframe.contentWindow.postMessage({ type: 'load_parcel_into_drawn', geometry: geometry }, '*');
                                } else if (attempts > 0) {
                                    setTimeout(() => trySend(attempts - 1), 300);
                                }
                            }
                            trySend(8);
                        } catch(e) {
                            console.warn('zoom geometry parse error:', e);
                        }
                    }
                    """
                )

                btn_update.click(
                    fn=update_parcel_dropdown,
                    inputs=[edit_id, e_name, e_resp, e_email, e_veg, e_infra, e_owner, e_concelho, e_area, edit_geom],
                    outputs=[edit_msg, parcel_dropdown, edit_map_loader, edit_id, e_name, e_resp, e_email, e_veg, e_infra, e_owner, e_concelho, e_area, edit_geom],
                    js="""
                    (pid, name, resp, email, veg, infra, owner, concelho, area, geom) => {
                        const geomEl = document.getElementById('geometry-input-edit');
                        const domGeom = geomEl ? (geomEl.querySelector('textarea') || {}).value : null;
                        const finalGeom = (domGeom && domGeom.trim()) ? domGeom : geom;
                        const areaEl = document.getElementById('edit-area-ha-input');
                        const domArea = areaEl ? (areaEl.querySelector('textarea') || areaEl.querySelector('input') || {}).value : null;
                        const finalArea = (domArea && domArea.trim()) ? domArea : area;
                        const concelhoEl = document.getElementById('edit-concelho-input');
                        const domConcelho = concelhoEl ? (concelhoEl.querySelector('textarea') || concelhoEl.querySelector('input') || {}).value : null;
                        const finalConcelho = (domConcelho && domConcelho.trim()) ? domConcelho : concelho;
                        return [pid, name, resp, email, veg, infra, owner, finalConcelho, finalArea, finalGeom];
                    }
                    """
                ).then(
                    fn=None,
                    inputs=None,
                    outputs=None,
                    js="""
                    () => {
                        const iframe = document.getElementById('edit-map-iframe');
                        if (iframe && iframe.contentWindow) {
                            // Clear drawn items then reset view to all parcels
                            iframe.contentWindow.postMessage({ type: 'clear_drawn_items' }, '*');
                            setTimeout(() => {
                                iframe.contentWindow.postMessage({ type: 'reset_view' }, '*');
                            }, 200);
                        }
                    }
                    """
                )

                btn_delete.click(
                    fn=delete_parcel_dropdown,
                    inputs=[edit_id],
                    outputs=[edit_msg, parcel_dropdown, edit_map_loader]
                ) 

            with gr.Tab("OPERACIONAIS"):
                # ── Stats bar ────────────────────────────────────────────
                op_stats_bar = gr.Markdown(value=get_operatives_stats())

                gr.Markdown("---")

                with gr.Row():
                    # ── Left: search + table ─────────────────────────────
                    with gr.Column(scale=3):
                        with gr.Row():
                            op_search = gr.Textbox(
                                label="🔍 Pesquisar",
                                placeholder="Nome, email, NIF ou telefone...",
                                scale=3
                            )
                            op_cert_filter = gr.Dropdown(
                                ["Todos", "Observador", "Operacional Queima", "Tecnico Fogo Controlado", "Outro"],
                                label="Certificação",
                                value="Todos",
                                scale=2
                            )

                        op_table = gr.Dataframe(
                            headers=["ID", "Nome", "NIF", "Email", "Telefone", "Certificação", "Nº Ações"],
                            value=get_filtered_operatives("", ""),
                            label="Operacionais — clique numa linha para carregar",
                            interactive=False
                        )

                    # ── Right: form ──────────────────────────────────────
                    with gr.Column(scale=2):
                        op_form_title = gr.Markdown("### ➕ Novo Operacional")
                        op_id_state = gr.State(value=None)

                        op_name  = gr.Textbox(label="Nome Completo *")
                        op_nif   = gr.Textbox(label="NIF", placeholder="Ex: 123456789")
                        op_email = gr.Textbox(label="Email *", placeholder="nome@example.com")
                        op_phone = gr.Textbox(label="Telefone", placeholder="+351 912 345 678")
                        op_cert  = gr.Dropdown(
                            ["Observador", "Operacional Queima", "Tecnico Fogo Controlado", "Outro"],
                            label="Certificação *",
                            value="Observador"
                        )
                        op_notes = gr.Textbox(label="Notas / Observações", lines=3, placeholder="Informação adicional...")

                        with gr.Row():
                            op_save_btn   = gr.Button("💾 Guardar", variant="primary")
                            op_clear_btn  = gr.Button("➕ Novo", variant="secondary")
                            op_delete_btn = gr.Button("🗑️ Eliminar", variant="stop")

                        op_msg = gr.Textbox(label="Estado", interactive=False)

                # ── Wire live search ─────────────────────────────────────
                def _filter(search, cert):
                    return get_filtered_operatives(search, cert)

                op_search.change(_filter, inputs=[op_search, op_cert_filter], outputs=[op_table])
                op_cert_filter.change(_filter, inputs=[op_search, op_cert_filter], outputs=[op_table])

                # ── Load operative into form on table row click ──────────
                def _load_row(evt: gr.SelectData, tbl):
                    try:
                        row = tbl.values.tolist()[evt.index[0]] if hasattr(tbl, 'values') else tbl[evt.index[0]]
                        op_id = int(row[0])
                        _, name, nif, email, phone, cert, notes = load_operative_for_edit(op_id)
                        return op_id, f"### ✏️ Editar — {name}", name, nif, email, phone, cert, notes, ""
                    except Exception as e:
                        return None, "### ➕ Novo Operacional", "", "", "", "", "Observador", "", f"Erro: {e}"

                op_table.select(
                    _load_row,
                    inputs=[op_table],
                    outputs=[op_id_state, op_form_title, op_name, op_nif, op_email, op_phone, op_cert, op_notes, op_msg]
                )

                # ── Clear / New button ───────────────────────────────────
                op_clear_btn.click(
                    lambda: (None, "### ➕ Novo Operacional", "", "", "", "", "Observador", "", ""),
                    outputs=[op_id_state, op_form_title, op_name, op_nif, op_email, op_phone, op_cert, op_notes, op_msg]
                )

                # ── Save (create or update) ──────────────────────────────
                op_save_btn.click(
                    save_operative,
                    inputs=[op_id_state, op_name, op_nif, op_email, op_phone, op_cert, op_notes],
                    outputs=[op_msg, op_table, op_stats_bar]
                ).then(
                    lambda: (None, "### ➕ Novo Operacional", "", "", "", "", "Observador", ""),
                    outputs=[op_id_state, op_form_title, op_name, op_nif, op_email, op_phone, op_cert, op_notes]
                )

                # ── Delete ───────────────────────────────────────────────
                op_delete_btn.click(
                    delete_operative_new,
                    inputs=[op_id_state],
                    outputs=[op_msg, op_table, op_stats_bar]
                ).then(
                    lambda: (None, "### ➕ Novo Operacional", "", "", "", "", "Observador", ""),
                    outputs=[op_id_state, op_form_title, op_name, op_nif, op_email, op_phone, op_cert, op_notes]
                )

            with gr.Tab("PRÉ-PLANO"):
                gr.Markdown("## 📋 Pré-Planeamento de Ações de Queima")
                gr.Markdown("*Registo preliminar: parcela(s), responsável, data indicativa e fotos pré-fogo.*")

                with gr.Row():
                    with gr.Column(scale=3):
                        pp_table = gr.Dataframe(
                            headers=["ID", "Nome", "Responsável", "Data Indicativa"],
                            value=get_preplans_list(),
                            label="Pré-Planos — clique para editar",
                            column_widths=["50px", "40%", "30%", "20%"],
                            interactive=False
                        )
                        gr.Markdown("### 📷 Fotos pré-fogo")
                        pp_gallery = gr.Gallery(
                            value=[],
                            columns=3,
                            height=280,
                            object_fit="cover",
                            show_label=False,
                            allow_preview=True
                        )
                    with gr.Column(scale=2):
                        pp_form_title = gr.Markdown("### ➕ Novo Pré-Plano")
                        pp_id_state   = gr.State(value=None)

                        pp_name    = gr.Textbox(label="Nome / Designação *")
                        pp_resp    = gr.Textbox(label="Responsável", placeholder="Nome do responsável...")
                        pp_date    = gr.DateTime(label="Data Indicativa", include_time=False, type="string")
                        pp_parcels = gr.Dropdown(
                            choices=get_parcel_choices(),
                            label="Parcela(s) *",
                            multiselect=True
                        )
                        pp_notes  = gr.Textbox(label="Notas", lines=3)
                        pp_photos = gr.File(
                            label="Fotos da parcela (pré-fogo)",
                            file_count="multiple",
                            file_types=["image"]
                        )
                        with gr.Row():
                            pp_save_btn   = gr.Button("💾 Guardar", variant="primary")
                            pp_clear_btn  = gr.Button("➕ Novo", variant="secondary")
                            pp_delete_btn = gr.Button("🗑️ Eliminar", variant="stop")
                        pp_msg = gr.Textbox(label="Estado", interactive=False)

                def _pp_clear():
                    return None, "### ➕ Novo Pré-Plano", "", "", None, [], "", None, "", []

                def _pp_load_row(evt: gr.SelectData, tbl):
                    try:
                        row = tbl.values.tolist()[evt.index[0]] if hasattr(tbl, "values") else tbl[evt.index[0]]
                        pp_id, name, responsible, date, parcel_ids, notes, _, photos = load_preplan(int(row[0]))
                        return pp_id, f"### ✏️ Editar — {name}", name, responsible, date, parcel_ids, notes, None, "", photos
                    except Exception as e:
                        return None, "### ➕ Novo Pré-Plano", "", "", None, [], "", None, f"Erro: {e}", []

                _PP_OUTPUTS = [pp_id_state, pp_form_title, pp_name, pp_resp, pp_date, pp_parcels, pp_notes, pp_photos, pp_msg, pp_gallery]

                pp_table.select(
                    _pp_load_row, inputs=[pp_table],
                    outputs=_PP_OUTPUTS
                )
                pp_clear_btn.click(_pp_clear, outputs=_PP_OUTPUTS)

                def _pp_save(pp_id, name, resp, date, parcels, notes, photos):
                    msg, tbl = save_preplan(pp_id, name, resp, date, parcels, notes, photos)
                    # Reload gallery — works for both create (new id) and update (same id)
                    # We need to re-fetch the id after create, so reload from list
                    from fire_actions.models import FireAction as FA
                    try:
                        latest = FA.objects.filter(status="Pre-Plano").order_by('-id').first()
                        gal_id = pp_id if pp_id else (latest.id if latest else None)
                        gallery = _get_photo_paths(gal_id) if gal_id else []
                    except Exception:
                        gallery = []
                    return msg, tbl, gallery

                pp_save_btn.click(
                    _pp_save,
                    inputs=[pp_id_state, pp_name, pp_resp, pp_date, pp_parcels, pp_notes, pp_photos],
                    outputs=[pp_msg, pp_table, pp_gallery]
                ).then(_pp_clear, outputs=_PP_OUTPUTS)

                pp_delete_btn.click(
                    delete_preplan,
                    inputs=[pp_id_state],
                    outputs=[pp_msg, pp_table]
                ).then(_pp_clear, outputs=_PP_OUTPUTS)

            with gr.Tab("PLANO DE QUEIMA"):
                gr.Markdown("## 🔥 Plano Operacional de Queima")

                with gr.Row():
                    with gr.Column(scale=3):
                        bp_table = gr.Dataframe(
                            headers=["ID", "Pré-Plano", "Data Execução", "Parcelas", "Nº Homens"],
                            value=get_burning_plans_list(),
                            label="Planos de Queima — clique para carregar",
                            interactive=False
                        )
                    with gr.Column(scale=3):
                        bp_form_title = gr.Markdown("### ➕ Novo Plano de Queima")
                        bp_id_state   = gr.State(value=None)

                        with gr.Row():
                            bp_preplan = gr.Dropdown(
                                choices=get_preplan_choices(),
                                label="Pré-Plano associado *",
                                value=None,
                                scale=4,
                                allow_custom_value=True
                            )
                            bp_refresh_pp_btn = gr.Button("🔄", scale=1, min_width=60)
                        bp_date = gr.DateTime(label="Data de Execução *", include_time=False, type="string")

                        gr.Markdown("### 👥 Equipa")
                        bp_operatives = gr.Dropdown(
                            choices=get_operative_choices(),
                            label="Operacionais presentes",
                            multiselect=True,
                            allow_custom_value=True
                        )
                        bp_num_men = gr.Number(label="Nº de Homens", minimum=0, precision=0)

                        gr.Markdown("#### 🚒 Veículos")
                        with gr.Row():
                            bp_vfci      = gr.Number(label="VFCI",  minimum=0, precision=0, value=0)
                            bp_vfcm      = gr.Number(label="VLCI",  minimum=0, precision=0, value=0)
                            bp_other_veh = gr.Textbox(label="Outros veículos", placeholder="Ex: 1 ATV")

                        gr.Markdown("### ⚠️ Problemas identificados")
                        bp_problems = gr.CheckboxGroup(choices=PROBLEMS_CHOICES, label="", value=[])

                        gr.Markdown("### 🌿 Humidade dos Combustíveis")
                        with gr.Row():
                            bp_fuel_sup = gr.Textbox(label="Superficial (%)",  placeholder="Ex: 12")
                            bp_fuel_mf  = gr.Textbox(label="Manta morta F (%)", placeholder="Ex: 18")
                            bp_fuel_mh  = gr.Textbox(label="Manta morta H (%)", placeholder="Ex: 25")

                        gr.Markdown("### 🌤️ Meteorologia")
                        bp_weather_state = gr.Dropdown(choices=WEATHER_STATE_CHOICES, label="Estado do tempo")
                        with gr.Row():
                            bp_wind_bft = gr.Textbox(label="Velocidade vento (Beaufort)", placeholder="Ex: 3")
                            bp_wind_kmh = gr.Textbox(label="Velocidade vento (km/h)",     placeholder="Ex: 19-28")
                            bp_wind_dir = gr.Dropdown(choices=WIND_DIR_CHOICES, label="Direcção do vento")
                        bp_fire_conduct = gr.Dropdown(choices=FIRE_CONDUCT_CHOICES, label="Condução do fogo")
                        bp_fire_conduct_other = gr.Textbox(label="Descrever condução (se 'Outro')", visible=False)

                        gr.Markdown("### 📊 Efeitos e Eficácia")
                        bp_effects      = gr.Textbox(label="Efeitos da queima",  lines=3, placeholder="Descreva os efeitos observados...")
                        bp_efficiency   = gr.Textbox(label="Eficácia da ação",   lines=2, placeholder="Avaliação da eficácia...")
                        bp_extra_notes  = gr.Textbox(label="Notas adicionais",   lines=2)

                        gr.Markdown("### 📷 Fotografias da Queima")
                        bp_photos = gr.File(
                            label="Adicionar fotos (jpg/png)",
                            file_count="multiple",
                            file_types=["image"]
                        )
                        bp_gallery = gr.Gallery(
                            value=[],
                            columns=3,
                            height=220,
                            object_fit="cover",
                            show_label=False,
                            allow_preview=True
                        )

                        with gr.Row():
                            bp_save_btn   = gr.Button("💾 Guardar Plano", variant="primary")
                            bp_clear_btn  = gr.Button("➕ Novo",           variant="secondary")
                            bp_delete_btn = gr.Button("🗑️ Eliminar",       variant="stop")
                        bp_report_btn  = gr.Button("📄 Gerar Relatório Word", variant="secondary")
                        bp_report_file = gr.File(label="Relatório gerado", visible=False, interactive=False)
                        bp_msg = gr.Textbox(label="Estado", interactive=False)

                # Show/hide "Outro" description field
                bp_fire_conduct.change(
                    lambda v: gr.update(visible="5 -" in (v or "")),
                    inputs=[bp_fire_conduct],
                    outputs=[bp_fire_conduct_other]
                )

                # Refresh pre-plan dropdown on button click
                def _bp_refresh_preplans():
                    return gr.update(choices=get_preplan_choices())

                bp_refresh_pp_btn.click(_bp_refresh_preplans, outputs=[bp_preplan])

                # Cross-tab: refresh bp_preplan whenever a pre-plan is saved or deleted
                pp_save_btn.click(_bp_refresh_preplans, outputs=[bp_preplan])
                pp_delete_btn.click(_bp_refresh_preplans, outputs=[bp_preplan])

                _BP_OUTPUTS = [
                    bp_id_state, bp_form_title,
                    bp_preplan, bp_date, bp_operatives, bp_num_men,
                    bp_vfci, bp_vfcm, bp_other_veh,
                    bp_problems, bp_fuel_sup, bp_fuel_mf, bp_fuel_mh,
                    bp_weather_state, bp_wind_bft, bp_wind_kmh, bp_wind_dir,
                    bp_fire_conduct, bp_fire_conduct_other,
                    bp_effects, bp_efficiency, bp_extra_notes,
                    bp_photos, bp_gallery, bp_msg
                ]

                def _bp_clear():
                    #       id_state, title,
                    #       preplan, date, operatives, num_men,
                    #       vfci, vfcm, other_veh,
                    #       problems, fuel_sup, fuel_mf, fuel_mh,
                    #       weather_state, wind_bft, wind_kmh, wind_dir,
                    #       fire_conduct, fire_conduct_other,
                    #       effects, efficiency, extra_notes, msg
                    return (None, "### ➕ Novo Plano de Queima",
                            gr.update(choices=get_preplan_choices(), value=None),
                            None, [], None, 0, 0, "",
                            [], "", "", "",
                            None, "", "", None, None, "",
                            "", "", "", None, [], "")

                def _bp_load_row(evt: gr.SelectData, tbl):
                    try:
                        row = tbl.values.tolist()[evt.index[0]] if hasattr(tbl, "values") else tbl[evt.index[0]]
                        vals = load_burning_plan(int(row[0]))
                        bp_id = vals[0]
                        title = f"### ✏️ Editar Plano #{bp_id}"
                        # vals: id, preplan_id, date, ops, problems, fuel_sup, fuel_mf, fuel_mh,
                        #        wx, wind_bft, wind_kmh, wind_dir, conduct, conduct_other,
                        #        num_men, vfci_str, vfcm_str, other_veh, effects, efficiency, notes
                        fresh_pp   = get_preplan_choices()
                        fresh_ops  = get_operative_choices()
                        return (bp_id, title,
                                gr.update(choices=fresh_pp,  value=vals[1]),
                                vals[2],
                                gr.update(choices=fresh_ops, value=vals[3]),
                                vals[14],
                                int(vals[15] or 0), int(vals[16] or 0), vals[17],
                                vals[4], vals[5], vals[6], vals[7],
                                vals[8], vals[9], vals[10], vals[11],
                                vals[12], vals[13],
                                vals[18], vals[19], vals[20],
                                None, _get_bp_photo_paths(vals[0]), "")
                    except Exception as e:
                        return _bp_clear()[:-1] + (f"Erro: {e}",)

                bp_table.select(_bp_load_row, inputs=[bp_table], outputs=_BP_OUTPUTS)
                bp_clear_btn.click(_bp_clear, outputs=_BP_OUTPUTS)

                def _do_save_bp(bp_id, preplan_id, exec_date,
                                ops, num_men, vfci, vfcm, other_veh,
                                problems, fuel_sup, fuel_mf, fuel_mh,
                                wx_state, wind_bft, wind_kmh, wind_dir,
                                fire_conduct, fire_conduct_other,
                                effects, efficiency, extra_notes, photos):
                    import json as _json
                    vehicles_json = _json.dumps({
                        "VFCI": int(vfci or 0),
                        "VLCI": int(vfcm or 0),
                        "Outro": other_veh or ""
                    })
                    return save_burning_plan(
                        bp_id, preplan_id, exec_date,
                        ops, problems,
                        fuel_sup, fuel_mf, fuel_mh,
                        wx_state, wind_bft, wind_kmh, wind_dir,
                        fire_conduct, fire_conduct_other,
                        num_men, vehicles_json,
                        effects, efficiency, extra_notes,
                        photos=photos
                    )

                bp_save_btn.click(
                    _do_save_bp,
                    inputs=[bp_id_state, bp_preplan, bp_date,
                            bp_operatives, bp_num_men, bp_vfci, bp_vfcm, bp_other_veh,
                            bp_problems, bp_fuel_sup, bp_fuel_mf, bp_fuel_mh,
                            bp_weather_state, bp_wind_bft, bp_wind_kmh, bp_wind_dir,
                            bp_fire_conduct, bp_fire_conduct_other,
                            bp_effects, bp_efficiency, bp_extra_notes, bp_photos],
                    outputs=[bp_msg, bp_table]
                ).then(_bp_clear, outputs=_BP_OUTPUTS)

                bp_delete_btn.click(
                    delete_burning_plan,
                    inputs=[bp_id_state],
                    outputs=[bp_msg, bp_table]
                ).then(_bp_clear, outputs=_BP_OUTPUTS)

                def _bp_generate_report(bp_id):
                    file_path, msg = generate_burning_plan_report(bp_id)
                    photos = _get_bp_photo_paths(bp_id) if bp_id else []
                    if file_path:
                        return gr.update(value=file_path, visible=True), photos, msg
                    return gr.update(visible=False), photos, msg

                bp_report_btn.click(
                    _bp_generate_report,
                    inputs=[bp_id_state],
                    outputs=[bp_report_file, bp_gallery, bp_msg]
                )

            with gr.Tab("METEOROLOGIA"):
                # Construir iframe com parcelas da BD injetadas no FiredPT
                gr.HTML(build_weather_map_iframe())

            hourly_location = gr.Textbox(
                visible=True,
                elem_id="hourly-location",
                elem_classes="hidden-textbox"
            )

    main_app.load(
            fn=get_parcels_geojson, 
            outputs=[geojson_data_hidden]
        ).then(
            fn=_bp_clear,
            outputs=_BP_OUTPUTS
        )
    return main_app

def create_login_interface():
    """Interface de login"""
    with gr.Blocks() as login_app:
        gr.Markdown("# 🔥 Fogo Bom Algarve")
        gr.Markdown("## Sistema de Gestão de Fogo Controlado")
        
        with gr.Tab("🔐 Login"):
            with gr.Column(elem_classes="login-container"):
                gr.Markdown("### Iniciar Sessão")
                login_username = gr.Textbox(label="Nome de Utilizador", placeholder="username")
                login_password = gr.Textbox(label="Password", type="password", placeholder="••••••••")
                login_btn = gr.Button("Entrar", variant="primary", size="lg")
                login_message = gr.Textbox(label="Estado", interactive=False)
        
        with gr.Tab("📝 Registo"):
            with gr.Column(elem_classes="login-container"):
                gr.Markdown("### Criar Nova Conta")
                reg_username = gr.Textbox(label="Nome de Utilizador *", placeholder="username")
                reg_email = gr.Textbox(label="Email *", placeholder="email@example.com")
                reg_password = gr.Textbox(label="Password *", type="password", placeholder="Mínimo 6 caracteres")
                reg_confirm = gr.Textbox(label="Confirmar Password *", type="password", placeholder="••••••••")
                reg_btn = gr.Button("Criar Conta", variant="primary", size="lg")
                reg_message = gr.Textbox(label="Estado", interactive=False)
                gr.Markdown("*Após criar a conta, use o tab Login para entrar*")
        
        # Register handler
        reg_btn.click(
            register_user,
            inputs=[reg_username, reg_email, reg_password, reg_confirm],
            outputs=reg_message
        )
    
    return login_app, login_btn, login_username, login_password, login_message
# ============= CUSTOM CSS FOR TAB STYLING =============
CUSTOM_CSS = """
/* Auto-fill boxes (concelho + area) */
.autofill-box textarea, .autofill-box input {
    background-color: #fef9ec !important;
    border: 1px solid #f5c842 !important;
    color: #7a5c00 !important;
}
.autofill-box label {
    color: #b07d00 !important;
    font-style: italic !important;
}

/* Tab buttons styling - Responsive design */
button[role="tab"] {
    font-weight: 600 !important;
    font-size: 13px !important;
    padding: 8px 16px !important;
    margin: 0 2px !important;
    border-radius: 8px 8px 0 0 !important;
    transition: all 0.3s ease !important;
    border: none !important;
    border-bottom: 3px solid transparent !important;
    background: linear-gradient(135deg, #f5f5f4 0%, #e7e5e4 100%) !important;
    color: #57534e !important;
    position: relative !important;
}

button[role="tab"]:hover {
    background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%) !important;
    color: #92400e !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 8px rgba(234, 88, 12, 0.15) !important;
}

button[role="tab"][aria-selected="true"] {
    background: linear-gradient(135deg, #fb923c 0%, #ea580c 100%) !important;
    color: white !important;
    border-bottom: 3px solid #c2410c !important;
    box-shadow: 0 4px 12px rgba(234, 88, 12, 0.3) !important;
    font-weight: 700 !important;
}

/* Tab list container */
div[role="tablist"] {
    border-bottom: 2px solid #e7e5e4 !important;
    margin-bottom: 20px !important;
    gap: 2px !important;
    padding: 0 4px !important;
}

/* Responsive adjustments for larger screens */
@media (min-width: 1024px) {
    button[role="tab"] {
        font-size: 14px !important;
        padding: 10px 20px !important;
        margin: 0 3px !important;
    }
    
    div[role="tablist"] {
        gap: 3px !important;
        padding: 0 8px !important;
    }
}

@media (min-width: 1440px) {
    button[role="tab"] {
        font-size: 15px !important;
        padding: 12px 24px !important;
        margin: 0 4px !important;
    }
    
    div[role="tablist"] {
        gap: 4px !important;
        padding: 0 10px !important;
    }
}

/* Responsive adjustments for smaller screens */
@media (max-width: 768px) {
    button[role="tab"] {
        font-size: 12px !important;
        padding: 6px 12px !important;
        margin: 0 1px !important;
    }
    
    div[role="tablist"] {
        gap: 1px !important;
        padding: 0 2px !important;
    }
}

@media (max-width: 480px) {
    button[role="tab"] {
        font-size: 11px !important;
        padding: 5px 8px !important;
        margin: 0 1px !important;
    }
    
    div[role="tablist"] {
        gap: 1px !important;
        padding: 0 !important;
    }
}

/* Login container styling */
.login-container {
    max-width: 400px !important;
    margin: 40px auto !important;
    padding: 30px !important;
    border-radius: 12px !important;
    background: white !important;
    box-shadow: 0 8px 24px rgba(0,0,0,0.1) !important;
}
"""
# ============= MAIN APP =============
with gr.Blocks(title="Fogo Bom Algarve - Sistema de Gestão") as app:
    # Criar ambas as interfaces
    
    main_interface= create_main_interface()
    
    # Mostrar login por padrão

    main_interface.visible = True    
   
   

app.launch(
    server_name="0.0.0.0",
    server_port=7860,
    auth=None,
    allowed_paths=["./images", ".", str(MEDIA_DIR), str(BP_MEDIA_DIR)],
    share=False,
    theme=gr.themes.Default(primary_hue="stone", secondary_hue="orange"),
    css=CUSTOM_CSS,
    head=HEAD_JS)